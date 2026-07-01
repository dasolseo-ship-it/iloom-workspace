import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
section = doc.sections[0]
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

RED  = (0xC8, 0x0A, 0x1E)
GRAY = 'F0F0F0'
LGRAY= 'F8F8F8'

# ── helpers ────────────────────────────────────────────────
def rf(run, sz=9, bold=False, col=None):
    run.font.name = 'Malgun Gothic'
    run.font.size = Pt(sz)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    if col: run.font.color.rgb = RGBColor(*col)

def cw(cell, txt, bold=False, sz=9, col=None,
        align=WD_ALIGN_PARAGRAPH.LEFT, valign=WD_ALIGN_VERTICAL.CENTER):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    cell.vertical_alignment = valign
    if txt:
        r = p.add_run(txt); rf(r, sz, bold, col)
    return cell

def cbg(cell, hex_):
    pr = cell._tc.get_or_add_tcPr()
    s  = OxmlElement('w:shd')
    s.set(qn('w:val'),'clear'); s.set(qn('w:color'),'auto'); s.set(qn('w:fill'), hex_)
    pr.append(s)

def tb(table, c='BBBBBB'):
    bs = OxmlElement('w:tblBorders')
    for side in ['top','left','bottom','right','insideH','insideV']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'4')
        b.set(qn('w:space'),'0');    b.set(qn('w:color'), c)
        bs.append(b)
    table._tbl.tblPr.append(bs)

def secttitle(txt):
    p = doc.add_paragraph(); r = p.add_run(txt); rf(r, 10, True, RED)

def note(txt):
    p = doc.add_paragraph(); r = p.add_run(txt); rf(r, 8, col=(0x88,0x88,0x88))

def sub(txt):
    p = doc.add_paragraph(); r = p.add_run(txt); rf(r, 9, True)

def bullet(txt):
    p = doc.add_paragraph(); r = p.add_run(txt); rf(r)

def hdr_row(table, cols):
    for i,h in enumerate(cols):
        cw(table.rows[0].cells[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        cbg(table.rows[0].cells[i], GRAY)

# ── 제목 ───────────────────────────────────────────────────
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
rf(p.add_run('품의서 (자유양식)'), 18, True)
doc.add_paragraph()

# ── 결재란 ─────────────────────────────────────────────────
ap = doc.add_table(rows=3, cols=6)
ap.alignment = WD_TABLE_ALIGNMENT.RIGHT; tb(ap)
for i,h in enumerate(['작성','검토 I','검토 II','검토 III','검토 IV','승인']):
    cw(ap.rows[0].cells[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    cbg(ap.rows[0].cells[i], GRAY)
for row in ap.rows[1:]:
    for c in row.cells: c.add_paragraph(); c.add_paragraph()
doc.add_paragraph()

# ── 문서정보 ────────────────────────────────────────────────
inf = doc.add_table(rows=3, cols=4); tb(inf)
idata = [
    ('문서번호','일룸-품의26-XX-XXXXX','작성일','2026-06-30'),
    ('작성부서','일룸 > 일룸사업부 > 영업개발부문 > 대리점파트','작성자','서다솔'),
    ('제목','전주 서신더샵비발디 입주 공략의 건(중화산4점)(입주박람회)','열람권한','부서'),
]
IW = [2.5, 9.5, 2.0, 2.0]
for ri,rd in enumerate(idata):
    for ci,t in enumerate(rd):
        lbl = (ci%2==0)
        cw(inf.rows[ri].cells[ci], t, bold=lbl,
           align=WD_ALIGN_PARAGRAPH.CENTER if lbl else WD_ALIGN_PARAGRAPH.LEFT)
        if lbl: cbg(inf.rows[ri].cells[ci], GRAY)
for row in inf.rows:
    for i,c in enumerate(row.cells): c.width = Cm(IW[i])
doc.add_paragraph()

for txt in [
    '아래와 같이 전주 서신더샵비발디 입주아파트를 공략하고자 하오니 검토 후 재가 바랍니다.',
    '----- 아   래 -----',
]:
    p = doc.add_paragraph(txt); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; rf(p.runs[0])
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════
# ※ 요약 테이블
# ═══════════════════════════════════════════════════════════
secttitle('※ 본 입주공략 관련 요약')

# 9 rows × 3 cols
# 0: 헤더  1: 공략대상지  2: 세대/시점  3: 목표
# 4~6: 운영프로모션(좌셀 merge) / 4=상시박람회, 5=특별, 6=홍보마케팅
# 7: 공략채산
sm = doc.add_table(rows=8, cols=3); tb(sm)
SW = [3.5, 9.5, 3.0]

hdr_row(sm, ['구분','내용','비고'])

def sm_row(r, label, content, note_='', lbl_bg=True):
    cw(sm.rows[r].cells[0], label, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    if lbl_bg: cbg(sm.rows[r].cells[0], GRAY)
    cw(sm.rows[r].cells[1], content)
    cw(sm.rows[r].cells[2], note_)

sm_row(1, '공략대상지',       '전주 서신더샵비발디',              '중화산4점 단독 공략')
sm_row(2, '입주 세대 수\n및 시점', '1,225세대 (일반분양) / 2026년 11월 입주', '일반분양 1,225 / 조합원 574\n임대 96 / 보류 19 / 전체 1,914')
sm_row(3, '공략지 목표',
       '매출 약 1.0억 / 수주 약 49건 (수주율 4%) / 객단가 약 200만원',
       'VAT별도')

# 운영프로모션 좌셀 merge
sm.rows[4].cells[0].merge(sm.rows[6].cells[0])
cw(sm.rows[4].cells[0], '운영\n프로모션', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
cbg(sm.rows[4].cells[0], GRAY)

# Row4 — 상시/박람회: nested mini-table
c4 = sm.rows[4].cells[1]; c4.text = ''
p4 = c4.paragraphs[0]; rf(p4.add_run('상시 / 박람회'), True)
ht = c4.add_table(rows=6, cols=4); tb(ht, 'CCCCCC')
혜택 = [
    ('구매금액대','기존혜택(상시)','추가혜택(박람회전용)','총 혜택금액'),
    ('200만원',  '-',           '10만원(구매금액 5%)', '10만원'),
    ('300만원',  '15만원',       '15만원(구매금액 5%)', '30만원'),
    ('500만원',  '30만원',       '25만원(구매금액 5%)', '55만원'),
    ('700만원',  '40만원',       '35만원(구매금액 5%)', '75만원'),
    ('1,000만원','80만원',       '50만원(구매금액 5%)','130만원'),
]
for ri,rd in enumerate(혜택):
    for ci,t in enumerate(rd):
        sc = ht.rows[ri].cells[ci]
        cw(sc, t, bold=(ri==0), sz=8, align=WD_ALIGN_PARAGRAPH.CENTER)
        if ri==0: cbg(sc, LGRAY)

# Row4 비고셀
cw(sm.rows[4].cells[2],
   '※ 200만원 구간 혜택 및\n추가혜택(박람회전용)\n본사:대리점 = 5:5 부담')

# Row5 — 특별 프로모션
c5 = sm.rows[5].cells[1]; c5.text = ''
rf(c5.paragraphs[0].add_run('특별 (특정 구매 프로모션)'), bold=True)
for ln in ['(지정1) 지정세트 포함, 총 구매금액 300만원 이상 → 10만원 추가 혜택',
           '(지정2) 지정세트 포함, 총 구매금액 500만원 이상 → 20만원 추가 혜택']:
    p5 = c5.add_paragraph(); rf(p5.add_run(ln))

c5b = sm.rows[5].cells[2]; c5b.text = ''
for ln in ['지정세트 4종',
           '① 안방: 옷장 3통 이상',
           '② 주방: 식탁 4인+의자 4개 이상',
           '③ 거실: 소파 3인 이상',
           '④ 자녀방: SS침대+매트리스+옷장 1통 이상']:
    if not c5b.paragraphs[0].runs:
        rf(c5b.paragraphs[0].add_run(ln), sz=8)
    else:
        rf(c5b.add_paragraph().add_run(ln), sz=8)

# Row6 — 홍보마케팅
c6 = sm.rows[6].cells[1]; c6.text = ''
for ln in ['- 입주카페 주 1회 이상 게시글 작성','- 박람회 당일 계약후기 바이럴']:
    if not c6.paragraphs[0].runs: rf(c6.paragraphs[0].add_run(ln))
    else: rf(c6.add_paragraph().add_run(ln))

# Row7 — 공략채산
cw(sm.rows[7].cells[0], '공략 채산', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER); cbg(sm.rows[7].cells[0], GRAY)
cw(sm.rows[7].cells[1],
   '매출 약 1.0억 / *공략비용 약 7,600,000원 / 채산이익률 약 15.2% 예상')
cw(sm.rows[7].cells[2], '*공략비용: 프로모션비용 포함\n(추정값 — 채산 섹션 참조)')

for row in sm.rows:
    for i,c in enumerate(row.cells): c.width = Cm(SW[i])

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════
# 1. 배경 및 목적
# ═══════════════════════════════════════════════════════════
secttitle('1. 배경 및 목적')
for t in [
    '1) 전주 서신더샵비발디(1,225세대, 2026년 11월 입주 예정) 입주 수요를 공략하여 일룸 중화산4점 매출 확대 도모',
    '2) 라이프스타일 가구 브랜드 일룸 인지도 제고 및 지역 고객 접점 강화를 위한 입주박람회 단독 참가',
]: bullet(t)
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════
# 2. 공략대상
# ═══════════════════════════════════════════════════════════
secttitle('2. 공략대상')
note('(단위: 세대)')

# 4열: 대구분(col0) | 소구분(col1) | 내용(col2) | 비고(col3)
# rows 8: header + 4(입주아파트) + 상권정보 + 공략대상지 + 공략매장
tg = doc.add_table(rows=8, cols=4); tb(tg)

# header: col0+col1 merge → "구분"
tg.rows[0].cells[0].merge(tg.rows[0].cells[1])
cw(tg.rows[0].cells[0], '구분', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER); cbg(tg.rows[0].cells[0], GRAY)
cw(tg.rows[0].cells[2], '내용', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER); cbg(tg.rows[0].cells[2], GRAY)
cw(tg.rows[0].cells[3], '비고', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER); cbg(tg.rows[0].cells[3], GRAY)

# 입주아파트 — col0 vertical merge rows 1-4
tg.rows[1].cells[0].merge(tg.rows[4].cells[0])
cw(tg.rows[1].cells[0], '입주\n아파트', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER); cbg(tg.rows[1].cells[0], GRAY)

apt_data = [
    ('아파트명',      '전주 서신더샵비발디 (약명: 서신더샵)',
                      '시공: 포스코이앤씨 + 에이치엘디앤아이한라 컨소시엄'),
    ('위치',          '전북특별자치도 전주시 완산구 서신동 40-4번지 일대',
                      ''),
    ('세대수',        '일반분양 1,225세대 / 전체 1,914세대',
                      '조합원 574 / 임대 96 / 보류 19 포함'),
    ('세대구성',      '59㎡ 177세대  |  73㎡ 261세대  |  84㎡ 710세대  |  121㎡ 77세대',
                      '84㎡ 비중 57.9%'),
]
for ri, (label, content, bigo) in enumerate(apt_data):
    r_idx = ri + 1
    cw(tg.rows[r_idx].cells[1], label, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER); cbg(tg.rows[r_idx].cells[1], LGRAY)
    cw(tg.rows[r_idx].cells[2], content)
    cw(tg.rows[r_idx].cells[3], bigo)

# 상권 정보 — col0+col1 merge
tg.rows[5].cells[0].merge(tg.rows[5].cells[1])
cw(tg.rows[5].cells[0], '상권 정보', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER); cbg(tg.rows[5].cells[0], GRAY)
c5 = tg.rows[5].cells[2]; c5.text = ''
for i, ln in enumerate([
    '① 인프라',
    '  - 백제대로·전룡로 인접 / 호남고속도로 서전주IC 편리',
    '  - 롯데백화점·이마트 등 대형 유통시설 인근',
    '  - SRT 전라선 전주역 인근 (수서역 1시간 40분대)',
    '  - 전북대병원·원광대 전주한방병원 등 의료시설 인근',
    '  - 전주천·삼천 하천변 산책로, 안터공원 등 녹지 풍부',
    '② 학군',
    '  - 서신초·서신중·한일고 인근 / 전주시립 서신도서관',
]):
    p = c5.paragraphs[0] if i == 0 else c5.add_paragraph()
    r = p.add_run(ln); rf(r)

# 공략대상지 정보 — col0+col1 merge
tg.rows[6].cells[0].merge(tg.rows[6].cells[1])
cw(tg.rows[6].cells[0], '공략대상지\n정보', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER); cbg(tg.rows[6].cells[0], GRAY)
c6 = tg.rows[6].cells[2]; c6.text = ''
for i, ln in enumerate([
    '- 지상 10~20층 28개동, 총 1,914세대 (일반분양 1,225세대)',
    '- 84㎡ 타입이 일반분양의 58% → 옷장·붙박이장 수요 집중 예상',
    '- 일반분양 붙박이장 유상옵션 여부 확인 필요 (공략 품목 조정)',
    '- 소파·식탁 추가혜택으로 리빙 매출 확대 및 객단가 상승 유도',
    '- 입주 예정월: 2026년 11월 / 박람회 운영: 사장님 단독 응대',
]):
    p = c6.paragraphs[0] if i == 0 else c6.add_paragraph()
    r = p.add_run(ln); rf(r)

# 공략 매장 — col0+col1 merge
tg.rows[7].cells[0].merge(tg.rows[7].cells[1])
cw(tg.rows[7].cells[0], '공략 매장', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER); cbg(tg.rows[7].cells[0], GRAY)
cw(tg.rows[7].cells[2], '일룸 중화산4점')
cw(tg.rows[7].cells[3], '단독 공략')

# 열 너비
for row in tg.rows:
    row.cells[0].width = Cm(2.0)
    row.cells[1].width = Cm(2.5)
    row.cells[2].width = Cm(10.0)
    row.cells[3].width = Cm(2.5)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════
# 3. 홍보방안
# ═══════════════════════════════════════════════════════════
secttitle('3. 홍보방안')
sub('1) 입주카페')
for t in ['- 주 1회 이상 게시글 작성 (매장안내 / 제품소개 / 공간제안 / 프로모션 홍보 등)',
          '- 박람회 당일 계약후기 바이럴']:
    bullet(t)

doc.add_paragraph()
sub('2) 입주박람회')
note('(단위: 원, VAT별도)')

ct = doc.add_table(rows=6, cols=7); tb(ct)
hdr_row(ct, ['구분','총 비용','본사','분담률','대리점','분담률','비고'])
cost_data = [
    ('참가비',      3_500_000, 1_750_000, '50%', 1_750_000, '50%', ''),
    ('사인물 제작',   200_000,   100_000, '50%',   100_000, '50%', 'LED 배너 시안 / 예가'),
    ('리플릿 제작',   200_000,   100_000, '50%',   100_000, '50%', '3단 접지 / 예가'),
    ('프로모션비용', 6_500_000, 5_750_000, '88%',   750_000, '12%',
     '기존혜택(목표매출 5%) 100% 본사\n특별프로모션(세트구매 5%) 5:5\n세트구매비율 30% 가정'),
    ('합계',       10_400_000, 7_700_000, '74%', 2_700_000, '26%', ''),
]
for ri,rd in enumerate(cost_data):
    row = ct.rows[ri+1]
    is_total = (ri == len(cost_data)-1)
    cw(row.cells[0], rd[0], bold=is_total)
    for ci,val in enumerate([rd[1],rd[2],'',rd[4],'']):
        pass
    cw(row.cells[1], f'{rd[1]:,}', bold=is_total, align=WD_ALIGN_PARAGRAPH.CENTER)
    cw(row.cells[2], f'{rd[2]:,}', bold=is_total, align=WD_ALIGN_PARAGRAPH.CENTER)
    cw(row.cells[3], rd[3], bold=is_total, align=WD_ALIGN_PARAGRAPH.CENTER)
    cw(row.cells[4], f'{rd[4]:,}', bold=is_total, align=WD_ALIGN_PARAGRAPH.CENTER)
    cw(row.cells[5], rd[5], bold=is_total, align=WD_ALIGN_PARAGRAPH.CENTER)
    cw(row.cells[6], rd[6], sz=8)
    if is_total:
        for c in row.cells: cbg(c, LGRAY)

CW = [3.5, 2.0, 2.0, 1.5, 2.0, 1.5, 3.5]
for row in ct.rows:
    for i,c in enumerate(row.cells): c.width = Cm(CW[i])

note('※ 박람회 운전장비 및 퀵 비용은 지원하지 아니함')
note('※ 전시품 별도 이동 없음 (현장 사장님 단독 응대)')
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════
# 4. 프로모션
# ═══════════════════════════════════════════════════════════
secttitle('4. 프로모션')
sub('1) 상시입주 / 입주박람회 현장 프로모션')
note('(금액단위: VAT포함, 혜택 금액 부가세 없음, 이하동일)')

pt = doc.add_table(rows=7, cols=4); tb(pt)
hdr_row(pt, ['구매금액대 구분','기존혜택(상시)','추가혜택(박람회전용)','총 혜택금액'])

# add 기간/대상 rows before data
pt2 = doc.add_table(rows=8, cols=4); tb(pt2)
p_data_header = [
    ('구분','내용','','비고'),
    ('기간','2026년 8월 중 (구체 일정 확인 필요)','','기간 내 등록된 수주건에 한해 적용'),
    ('대상 시리즈','일룸 전 제품','',''),
]

# simpler: just혜택 table
혜택2 = doc.add_table(rows=7, cols=5); tb(혜택2)
혜택2_hdr = ['구분','내용','','','비고']
# Actually keep it simple:
혜택3 = doc.add_table(rows=8, cols=4); tb(혜택3)
for i,h in enumerate(['구매금액대 구분','기존혜택(상시)','추가혜택(박람회전용,\n지정세트 조건 충족 시)','총 혜택금액']):
    cw(혜택3.rows[0].cells[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER); cbg(혜택3.rows[0].cells[i], GRAY)

혜택3_data = [
    ('200만원',  '  -',    '10만원 (5%)', '10만원'),
    ('300만원',  '15만원', '15만원 (5%)', '30만원'),
    ('500만원',  '30만원', '25만원 (5%)', '55만원'),
    ('700만원',  '40만원', '35만원 (5%)', '75만원'),
    ('1,000만원','80만원', '50만원 (5%)','130만원'),
]
# 비고: 마지막 두 행에 병합 비고
for ri,rd in enumerate(혜택3_data):
    row = 혜택3.rows[ri+1]
    for ci,t in enumerate(rd):
        cw(row.cells[ci], t, align=WD_ALIGN_PARAGRAPH.CENTER)

# merge 비고셀 in 혜택3 — add a separate note
note('※ 200만원 구간 추가혜택 적용 여부는 채산 검토 후 대리점과 협의 예정')
note('※ 200만원 구간 혜택 및 추가혜택(박람회전용)의 경우 본사:대리점 = 5:5 부담')

doc.add_paragraph()
sub('2) 전주 서신더샵비발디 특별 프로모션')

sp = doc.add_table(rows=4, cols=3); tb(sp)
hdr_row(sp, ['프로모션명','특정 구매 프로모션','비고'])
sp_data = [
    ('기간',
     '2026년 7월 11일(토) ~ 7월 12일(일) / 2일간',
     '기간 내 수주 등록건에 한해 적용'),
    ('내용',
     '(지정1) 지정세트를 포함하여 총 구매금액 300만원 이상 시 10만원 추가 혜택\n'
     '(지정2) 지정세트를 포함하여 총 구매금액 500만원 이상 시 20만원 추가 혜택\n'
     '※ 지정세트:\n'
     '  ① 안방 세트: 옷장 3통 이상\n'
     '  ② 주방 세트: 식탁 4인 + 의자 4개 이상\n'
     '  ③ 거실 세트: 소파 3인 이상\n'
     '  ④ 자녀방 세트: SS 침대프레임 + 매트리스 + 옷장 1통 이상',
     '특별 프로모션의 경우,\n본사:대리점 = 5:5 부담\n상시입주/입주박람회 프로모션만 중복 가능'),
    ('건명',
     '프로모션 건명 지정: 고객명 뒤 \'(지정1)\' or \'(지정2)\' 기입\n'
     '예시) 홍길동(입주)(서신더샵)(지정1)',
     ''),
]
for ri,rd in enumerate(sp_data):
    row = sp.rows[ri+1]
    cw(row.cells[0], rd[0], bold=True, align=WD_ALIGN_PARAGRAPH.CENTER); cbg(row.cells[0], GRAY)
    c1 = row.cells[1]; c1.text = ''
    for li,ln in enumerate(rd[1].split('\n')):
        if li==0: rf(c1.paragraphs[0].add_run(ln))
        else: rf(c1.add_paragraph().add_run(ln))
    c2 = row.cells[2]; c2.text = ''
    for li,ln in enumerate(rd[2].split('\n')):
        if li==0:
            if ln: rf(c2.paragraphs[0].add_run(ln))
        else:
            if ln: rf(c2.add_paragraph().add_run(ln))

doc.add_paragraph()
sub('3) 공통사항')
for t in [
    '- 납기한도: 2026년 11월 30일 (입주 후 1개월)',
    '- LG전자 제휴 혜택 가능',
    '- 쿠시노반짝(및 이후 브랜드 스팟성 프로모션 포함) / 프레임반값과 중복 불가',
    '- 고객 혜택금은 천원단위 절사 지급 기준',
    '- 납기 종료 후 익월말 포인트 일괄 지급 (단, 특별 프로모션(특정 구매)은 ERP 등록 없이 영업담당자 별도 정산)',
]: bullet(t)

doc.add_paragraph()
sub('4) 정산')
for t in [
    '- 입주 공략 비용: 2026년 7월 (참가비 / 사인물 / 리플릿)',
    '- 프로모션 비용: 2026년 12월 (납기한도: 2026년 11월 30일)',
    '- \'입주박람회 프로모션 합의서\' 체결 내용에 따르며 품의 최종 승인 후 전자계약 진행 예정',
]: bullet(t)
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════
# 5. 예상매출 및 손익
# ═══════════════════════════════════════════════════════════
secttitle('5. 예상매출 및 손익')
bullet('1) 예상매출: 총 매출 약 1.0억 / 수주 약 49건 (수주율 4%) / 객단가 약 200만원 (VAT별도)')
bullet('2) 예상채산')
note('(단위: 원, VAT별도)  ※ 이하 수치는 추정값 — 목표 매출 확정 후 재산정 필요')

pnl = doc.add_table(rows=14, cols=3); tb(pnl)
hdr_row(pnl, ['구분','입주공략','비고'])
pnl_data = [
    # (label, amount, pct, note, indent, bold)
    ('매출',                100_000_000, 100.0, '',                                         False, True),
    ('  매출원가',           48_400_000,  48.4, '25년 리테일사업팀 연평균',                    True, False),
    ('  판매수수료 일반BS B타입 위탁판매', 10_000_000, 10.0, '일반BS B타입 위탁판매수수료(10%)', True, False),
    ('  판매수수료 PG사',     1_900_000,   1.9, '',                                          True, False),
    ('  소계',              11_900_000,  11.9, '',                                          True, False),
    ('  물류시공비',         10_600_000,  10.6, '25년 리테일사업팀 연평균',                    True, False),
    ('변동비 합계',          70_900_000,  70.9, '',                                         False, True),
    ('공헌이익',             29_100_000,  29.1, '',                                         False, True),
    ('  영업외(간접배부)',    10_100_000,  10.1, '25년 리테일사업팀 연평균',                    True, False),
    ('  광선비/판촉비',       3_800_000,   3.8, '참가비/경품/사인물/프로모션',                 True, False),
    ('고정비 합계',          13_900_000,  13.9, '',                                         False, True),
    ('채산이익',             15_200_000,  15.2, '',                                         False, True),
]
for ri,rd in enumerate(pnl_data):
    row = pnl.rows[ri+1]
    label, amt, pct, note_, indent, bold = rd
    cw(row.cells[0], label, bold=bold); cbg(row.cells[0], GRAY if not indent else 'FFFFFF')
    # amount + pct in col1
    cell = row.cells[1]; cell.text = ''
    p_ = cell.paragraphs[0]; p_.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf(p_.add_run(f'{amt:,}'), bold=bold)
    # pct in same cell second paragraph
    p2 = cell.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf(p2.add_run(f'{pct:.1f}%'), bold=bold, sz=8)
    cw(row.cells[2], note_, sz=8)
    if bold: [cbg(row.cells[i], LGRAY) for i in range(3)]

for row in pnl.rows:
    row.cells[0].width = Cm(5.0)
    row.cells[1].width = Cm(5.0)
    row.cells[2].width = Cm(6.0)

note('※ 영업외간접배부 제외 시 채산이익 25.3%')
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════
# 6. 기타사항
# ═══════════════════════════════════════════════════════════
secttitle('6. 기타사항')
for t in [
    '1) 특정 기간에만 적용 가능한 프로모션으로 "일룸 중화산4점"에서 진행',
    '2) 아파트 약명: 서신더샵 (건명 기입 기준)',
    '3) 기존 보관 중인 롤업 배너(전군 보관) 활용 — 시안만 제작하여 중화산4점으로 발송',
    '4) 3단 접지리플릿은 디곳통 디자인 후 인성 인쇄 발주 예정',
]: bullet(t)
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════
# 7. 첨부파일
# ═══════════════════════════════════════════════════════════
secttitle('7. 첨부파일')
for t in [
    '1) 전주 서신더샵비발디 입주박람회 입점 견적서 (접수 후 첨부 예정)',
    '2) 전주 서신더샵비발디 프로모션 합의서_중화산4점',
]: bullet(t)
doc.add_paragraph()

p = doc.add_paragraph('끝.'); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
rf(p.runs[0], bold=True)

# 저장
OUT = r'20-operations/24-입주공략/전주서신더샵비발디/2026-08_전주서신더샵비발디_입주공략_품의_v6.docx'
doc.save(OUT)
print(f'저장: {OUT}')
