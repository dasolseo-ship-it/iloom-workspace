import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
sec = doc.sections[0]
sec.top_margin    = Cm(2.5)
sec.bottom_margin = Cm(2.5)
sec.left_margin   = Cm(3.0)
sec.right_margin  = Cm(3.0)

RED  = (0xC8, 0x0A, 0x1E)
GRAY = 'F0F0F0'
LGRAY= 'F8F8F8'

def rf(run, sz=10, bold=False, col=None):
    run.font.name = 'Malgun Gothic'
    run.font.size = Pt(sz)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    if col: run.font.color.rgb = RGBColor(*col)

def para(txt, sz=10, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, col=None, space_before=0, space_after=4):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if txt:
        r = p.add_run(txt); rf(r, sz, bold, col)
    return p

def tb(table, c='BBBBBB'):
    bs = OxmlElement('w:tblBorders')
    for side in ['top','left','bottom','right','insideH','insideV']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'4')
        b.set(qn('w:space'),'0');    b.set(qn('w:color'), c)
        bs.append(b)
    table._tbl.tblPr.append(bs)

def cw(cell, txt, bold=False, sz=10, col=None,
        align=WD_ALIGN_PARAGRAPH.LEFT, valign=WD_ALIGN_VERTICAL.CENTER):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    cell.vertical_alignment = valign
    if txt:
        r = p.add_run(txt); rf(r, sz, bold, col)
    return cell

def cbg(cell, hex_):
    pr = cell._tc.get_or_add_tcPr()
    s  = OxmlElement('w:shd')
    s.set(qn('w:val'),'clear'); s.set(qn('w:color'),'auto'); s.set(qn('w:fill'), hex_)
    pr.append(s)

def article(num, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(f'제{num}조 {title}'); rf(r, 10, True)

def note(txt):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(txt); rf(r, 9, col=(0x66,0x66,0x66))

# ── 제목 ───────────────────────────────────────────────────
para('전주 서신더샵비발디 입주박람회 프로모션 합의서',
     sz=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=12)

# ── 전문 ───────────────────────────────────────────────────
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(8)
r = p.add_run(
    '주식회사 일룸(이하 "공급업자")과 일룸 중화산4점(이하 "대리점")은 공급업자의 특정행사'
    '(이하 "입주박람회")와 관련하여 아래와 같이 합의한다.'
)
rf(r, 10)

para('- 아   래 -', sz=10, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=2, space_after=8)

# ── 제1조 ──────────────────────────────────────────────────
article('1', '목적')
para('이 합의는 공급업자가 입주박람회 프로모션을 대리점과 함께 진행함을 목적으로 한다.',
     space_after=2)

# ── 제2조 ──────────────────────────────────────────────────
article('2', '합의의 효력')
para('본 합의의 효력은 입주박람회 기간 동안만 적용되며, 입주박람회 프로모션과 관련하여서는 '
     '본 합의가 당사자 간 다른 합의에 우선하여 효력을 발생한다.', space_after=2)

# ── 제3조 ──────────────────────────────────────────────────
article('3', '프로모션 비용의 분담')
para('본 프로모션 기간 동안 판매되는 상품의 종류는 전 제품을 대상으로 한다.', space_after=2)
para('공급업자와 대리점은 입주박람회 기간 동안 지출하는 비용을 아래와 같이 분담하기로 한다.',
     space_after=4)

# 비용 분담 테이블
ct = doc.add_table(rows=6, cols=3); tb(ct)
ct.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = ct.rows[0]
for i,(h,w) in enumerate(zip(['구분','세부항목','분담 비율'],[Cm(2.5),Cm(9.0),Cm(4.5)])):
    cw(hdr.cells[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER); cbg(hdr.cells[i], GRAY)
    hdr.cells[i].width = w

cost_rows = [
    ('참가비',     '박람회 참가비 (3,500,000원, VAT별도)',          '대리점 / 본사 각 50%'),
    ('광고판촉비', '사인물 제작 (LED 배너 시안)',                   '대리점 / 본사 각 50%'),
    ('광고판촉비', '리플릿 제작 (3단 접지)',                        '대리점 / 본사 각 50%'),
    ('프로모션',   '기존혜택 (구매금액대별 상시입주혜택)',           '본사 100% 부담'),
    ('프로모션',   '특별 프로모션 (특정 구매 프로모션)\n'
                   '- (지정1) 지정세트 포함 300만원 이상 → 10만원 추가 혜택\n'
                   '- (지정2) 지정세트 포함 500만원 이상 → 20만원 추가 혜택', '대리점 / 본사 각 50%'),
]
for ri, (cat, detail, ratio) in enumerate(cost_rows):
    row = ct.rows[ri+1]
    cw(row.cells[0], cat, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    cbg(row.cells[0], LGRAY)
    c1 = row.cells[1]; c1.text = ''
    for li, ln in enumerate(detail.split('\n')):
        p_ = c1.paragraphs[0] if li==0 else c1.add_paragraph()
        rf(p_.add_run(ln), 10)
    cw(row.cells[2], ratio, align=WD_ALIGN_PARAGRAPH.CENTER)
for row in ct.rows:
    row.cells[0].width = Cm(2.5)
    row.cells[1].width = Cm(9.0)
    row.cells[2].width = Cm(4.5)

doc.add_paragraph()
note('※ 전산장비 렌탈 및 퀵서비스 비용은 지원항목 제외')
note('※ 박람회 운전장비 및 전시품 별도 이동 비용은 지원항목 제외')

# 특별 프로모션 세트 조건
para('특별 프로모션 세트 조건', sz=10, bold=True, space_before=6, space_after=2)
st = doc.add_table(rows=5, cols=2); tb(st)
st.alignment = WD_TABLE_ALIGNMENT.CENTER
for i,(h,w) in enumerate(zip(['공간','세트 조건'],[Cm(3.0),Cm(13.0)])):
    cw(st.rows[0].cells[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER); cbg(st.rows[0].cells[i], GRAY)
    st.rows[0].cells[i].width = w
set_rows = [
    ('안방 세트',   '옷장 3통 이상'),
    ('주방 세트',   '식탁 4인 + 의자 4개 이상'),
    ('거실 세트',   '소파 3인 이상'),
    ('자녀방 세트', 'SS 침대프레임 + 매트리스 + 옷장 1통 이상'),
]
for ri,(space,cond) in enumerate(set_rows):
    cw(st.rows[ri+1].cells[0], space, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER); cbg(st.rows[ri+1].cells[0], LGRAY)
    cw(st.rows[ri+1].cells[1], cond)
    st.rows[ri+1].cells[0].width = Cm(3.0)
    st.rows[ri+1].cells[1].width = Cm(13.0)

doc.add_paragraph()
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
r = p.add_run(
    '단, 공급업자는 대리점과 합의된 매출목표 100,000,000원(VAT별도)을 달성할 경우, '
    '참가비 및 특별 프로모션 비용 중 대리점 분담액의 50%에 해당하는 금액을 목표달성 리워드로 지급하며, '
    '해당 리워드는 납기한도 익월에 파트너 리워드 포인트로 지급한다.'
)
rf(r, 10)

# ── 제4조 ──────────────────────────────────────────────────
article('4', '프로모션 기간')
for ln in [
    '∙ 입주박람회 프로모션 기간: 2026년 7월 11일(토) ~ 7월 12일(일) / 2일간',
    '∙ 품목별 프로모션 납기한도: 2026년 11월 30일 (입주 후 1개월)',
    '단, 공급업자와 대리점의 상호 서면 합의에 따라 해당 기간을 변경할 수 있다.',
]:
    para(ln, space_before=0, space_after=2)

# ── 제5조 ──────────────────────────────────────────────────
article('5', '정산')
for ln in [
    '∙ 프로모션 비용을 제외한 비용(참가비·광고판촉비)은 대리점이 선 집행 후, 제3조 분담 비율에 따라 지원한다.',
    '∙ 프로모션 비용 분담은 대리점/본사 분담비율에 따라 정산하며, 본사가 선 집행 후 납기한도(2026년 11월 30일) 이후 대리점에 일괄 청구한다.',
]:
    para(ln, space_before=0, space_after=2)

# ── 제6조 ──────────────────────────────────────────────────
article('6', '분쟁해결 및 재판관할')
para(
    '본 합의의 해석이나 이행에 관하여 양 당사자 간에 의견차이 또는 분쟁이 발생하는 경우 양 당사자는 '
    '원만한 합의를 통해 해결함을 원칙으로 하며, 합의가 이루어지지 아니한 경우에는 서울중앙지방법원을 '
    '제1심의 전속적 합의관할 법원으로 하여 소송을 통해 분쟁을 해결한다.',
    space_after=2
)

# ── 제7조 ──────────────────────────────────────────────────
article('7', '기타')
para('본 합의의 내용은 양 당사자의 별도 서면 합의에 의해 변경할 수 있다.', space_after=2)

# ── 마무리 문구 ────────────────────────────────────────────
para(
    '상기 합의 내용을 확인, 증명하기 위하여 본 합의서 2통을 작성하여 공급업자와 대리점이 '
    '서명 또는 날인 후 각각 1부씩 보관한다.',
    space_before=10, space_after=6
)
para('2026년 　　월 　　일', sz=10, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

# ── 서명란 ─────────────────────────────────────────────────
sg = doc.add_table(rows=4, cols=2); tb(sg)
sg.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['공급업자', '대리점']
for i, h in enumerate(headers):
    cw(sg.rows[0].cells[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER); cbg(sg.rows[0].cells[i], GRAY)

sign_data = [
    ('상  호', '주식회사 일룸',              '일룸 중화산4점'),
    ('주  소', '서울시 송파구 오금로 311 (오금동)', '전북특별자치도 전주시 완산구 (기입 필요)'),
    ('대표자', '대표이사  정보은　　(인)',      '대표자  　　　　　(인)'),
]
for ri, (label, left, right) in enumerate(sign_data):
    row = sg.rows[ri+1]
    # col0
    c0 = row.cells[0]; c0.text = ''
    p_ = c0.paragraphs[0]
    rf(p_.add_run(f'[{label}]  '), 9, True, col=(0x66,0x66,0x66))
    rf(p_.add_run(left), 10)
    # col1
    c1 = row.cells[1]; c1.text = ''
    p_ = c1.paragraphs[0]
    rf(p_.add_run(f'[{label}]  '), 9, True, col=(0x66,0x66,0x66))
    rf(p_.add_run(right), 10)

for row in sg.rows:
    row.cells[0].width = Cm(8.0)
    row.cells[1].width = Cm(8.0)

# ── 저장 ───────────────────────────────────────────────────
OUT = r'20-operations/24-입주공략/전주서신더샵비발디/2026-08_서신더샵비발디_입주박람회_프로모션합의서_초안.docx'
doc.save(OUT)
print(f'저장: {OUT}')
