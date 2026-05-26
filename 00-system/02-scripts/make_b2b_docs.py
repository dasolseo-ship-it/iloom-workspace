# -*- coding: utf-8 -*-
"""
B2B 협조전 (롯데영등포 / 구세군작업장) Word 파일 생성
+ B2B 거래 물품 할인공급 요청서 날짜 수정본 Word 생성
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = r"C:\Users\FURSYS\Downloads"

# ── 헬퍼 ──────────────────────────────────────────
def set_font(run, name="맑은 고딕", size=10, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    # 한글 폰트 강제 지정
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), name)
    rPr.insert(0, rFonts)

def add_para(doc, text="", align=WD_ALIGN_PARAGRAPH.LEFT, bold=False,
             size=10, space_before=0, space_after=4, color=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        run = p.add_run(text)
        set_font(run, size=size, bold=bold, color=color)
    return p

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), kwargs.get(side, "single"))
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), kwargs.get("color", "AAAAAA"))
        tcBorders.append(border)
    tcPr.append(tcBorders)

def cell_text(cell, text, size=9, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, color=None):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)


# ══════════════════════════════════════════════════
#  협조전 Word 생성
# ══════════════════════════════════════════════════
def make_hyeopjojeon():
    doc = Document()

    # 여백 설정
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # ── 제목 ──
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(14)
    r = p_title.add_run("협    조    전")
    set_font(r, size=18, bold=True)

    # ── 문서 정보 표 ──
    info_table = doc.add_table(rows=4, cols=4)
    info_table.style = "Table Grid"
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    widths = [Cm(2.5), Cm(5.5), Cm(2.5), Cm(5.5)]
    for row in info_table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = widths[i]

    def info_row(r_idx, label1, val1, label2, val2):
        row = info_table.rows[r_idx]
        for i, (txt, bold, bg) in enumerate([
            (label1, True,  "F0F0F0"),
            (val1,   False, "FFFFFF"),
            (label2, True,  "F0F0F0"),
            (val2,   False, "FFFFFF"),
        ]):
            cell_text(row.cells[i], txt, bold=bold, size=9,
                      align=WD_ALIGN_PARAGRAPH.CENTER if (i % 2 == 0)
                            else WD_ALIGN_PARAGRAPH.LEFT)
            set_cell_bg(row.cells[i], bg)

    info_row(0, "문서번호", "",              "작성일자", "2026-05-26")
    info_row(1, "작성부서", "일룸사업부 > 영업개발부문 > 리테일사업팀", "작성자", "서다솔")
    info_row(2, "제    목",
             "일룸 롯데영등포점 B2B 할인율 적용 협조요청(구세군작업장)",
             "열람권한", "부서")
    # 4번째 행: 제목 셀 병합
    info_table.rows[2].cells[1].merge(info_table.rows[2].cells[2])

    doc.add_paragraph()  # 여백

    # ── 인사말 ──
    add_para(doc,
             "아래와 같이 B2B 할인율 적용을 요청하오니 업무 협조 부탁드립니다.",
             align=WD_ALIGN_PARAGRAPH.LEFT, size=10, space_after=6)
    add_para(doc, "----- 아   래 -----",
             align=WD_ALIGN_PARAGRAPH.CENTER, size=10, space_after=10)

    # ── 1. 개요 ──
    add_para(doc, "1. 개요", bold=True, size=11, space_after=4)

    items_gaeyoh = [
        "1) 대리점명 : 롯데영등포점",
        "2) 수요처명 : 구세군작업장 (한국장애인 고용공단 보조공학센터)",
        "3) 견적품목 : 비토 스터디 의자, 싯브레이크 캐스터",
        "4) 진행사유",
    ]
    for txt in items_gaeyoh:
        p = add_para(doc, txt, size=10, space_before=0, space_after=2)
        p.paragraph_format.left_indent = Cm(0.5)

    for txt in [
        "- 수요처에서 관련 증빙서류(세금계산서) 대리점 발행 요청",
        "- B2B 처리 프로세스에 따라 10% 할인 적용 후, 대리점에서 수요처와 계약 진행",
    ]:
        p = add_para(doc, txt, size=10, space_before=0, space_after=2)
        p.paragraph_format.left_indent = Cm(1.0)

    doc.add_paragraph()

    # ── 2. 요청사항 ──
    add_para(doc, "2. 요청사항", bold=True, size=11, space_after=4)

    p = add_para(doc, "1) 해당 수주 21% 할인율을 적용 요청", size=10, space_after=4)
    p.paragraph_format.left_indent = Cm(0.5)

    # 금액 표
    p_unit = add_para(doc, "(단위 : 원 / VAT 포함)",
                      align=WD_ALIGN_PARAGRAPH.RIGHT, size=9, space_after=2)

    tbl = doc.add_table(rows=3, cols=6)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    col_widths = [Cm(2.4), Cm(3.5), Cm(3.8), Cm(2.5), Cm(2.8), Cm(2.6)]
    for row in tbl.rows:
        for i, cell in enumerate(row.cells):
            cell.width = col_widths[i]

    headers = ["대리점명", "수주번호", "건명", "수주금액\n(A)", "총 할인금액\n(B=A×21%)", "총 판매금액\n(C=A-B)"]
    for i, h in enumerate(headers):
        cell_text(tbl.rows[0].cells[i], h, bold=True, size=9)
        set_cell_bg(tbl.rows[0].cells[i], "C80A1E")
        tbl.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    data_row = ["롯데영등포", "I20260513126100", "구세군작업장\n(보조공학센터)", "3,350,000", "703,500", "2,646,500"]
    for i, d in enumerate(data_row):
        cell_text(tbl.rows[1].cells[i], d, size=9)

    total_row = ["합  계", "", "", "3,350,000", "703,500", "2,646,500"]
    for i, d in enumerate(total_row):
        cell_text(tbl.rows[2].cells[i], d, bold=True, size=9)
        set_cell_bg(tbl.rows[2].cells[i], "F5F5F5")
        if i in (3, 4, 5):
            tbl.rows[2].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(200, 10, 30)

    doc.add_paragraph()

    # ── 3. 첨부 ──
    add_para(doc, "3. 첨부", bold=True, size=11, space_after=4)
    for txt in [
        "1) 최종견적(B2B)_롯데영등포점",
        "2) 수주내역_롯데영등포점",
        "3) 사업자등록증_구세군작업장",
        "4) B2B 거래요청서 날인본_롯데영등포점",
    ]:
        p = add_para(doc, txt, size=10, space_before=0, space_after=3)
        p.paragraph_format.left_indent = Cm(0.5)

    # 저장
    path = os.path.join(OUT_DIR, "협조전_롯데영등포_B2B_구세군작업장.docx")
    doc.save(path)
    print(f"✅ 협조전 저장: {path}")
    return path


# ══════════════════════════════════════════════════
#  B2B 거래 물품 할인공급 요청서 날짜 수정본 Word 생성
# ══════════════════════════════════════════════════
def make_b2b_request():
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.5)
        section.right_margin  = Cm(3.0)

    # 제목
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(20)
    r = p.add_run("B2B 거래 물품 할인공급 요청서")
    set_font(r, size=16, bold=True)

    doc.add_paragraph()

    def add_item(label, content, size=11):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
        r1 = p.add_run(f"{label}  ")
        set_font(r1, size=size, bold=True)
        r2 = p.add_run(content)
        set_font(r2, size=size)

    add_item("1.", "대리점명 : 일룸 롯데영등포점")
    doc.add_paragraph()
    add_item("2.", "B2B 거래사유 : 고객요청 – 구매처인 '한국장애인 고용공단 보조공학센터'에서\n        공급사인 일룸의 세금계산서 발행을 요청함")
    doc.add_paragraph()
    add_item("3.", "수주정보")

    # 수주정보 표
    p_unit = add_para(doc, "(VAT 포함)",
                      align=WD_ALIGN_PARAGRAPH.RIGHT, size=9, space_after=2)

    tbl = doc.add_table(rows=3, cols=7)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    col_w = [Cm(2.2), Cm(3.6), Cm(2.8), Cm(2.4), Cm(2.2), Cm(2.2), Cm(2.4)]
    for row in tbl.rows:
        for i, cell in enumerate(row.cells):
            cell.width = col_w[i]

    # 헤더 병합 (수주금액 등은 2행 구조)
    hdrs_top = ["대리점명", "수주번호", "수주건명", "확정납기", "수주금액", "총 할인금액", "총 판매금액"]
    hdrs_bot = ["", "", "", "", "(A)", "(B)", "(C=A-B)"]

    for i, h in enumerate(hdrs_top):
        cell_text(tbl.rows[0].cells[i], h, bold=True, size=9)
        set_cell_bg(tbl.rows[0].cells[i], "404040")
        tbl.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)

    data = ["롯데영등포", "I20260513126100", "구세군작업장", "2026-05-27", "3,350,000", "703,500", "2,646,500"]
    for i, d in enumerate(data):
        cell_text(tbl.rows[1].cells[i], d, size=9)

    total = ["합계", "", "", "", "3,350,000", "703,500", "2,646,500"]
    for i, d in enumerate(total):
        cell_text(tbl.rows[2].cells[i], d, bold=True, size=9)
        set_cell_bg(tbl.rows[2].cells[i], "F0F0F0")

    doc.add_paragraph()
    doc.add_paragraph()

    # 4. 요청 내용 상세
    add_item("4.", "요청 내용 상세")

    for txt in [
        "1) 상기 물품 할인공급 건은 기존 체결된 위탁계약과 별개의 건임을 고지합니다.",
        "2) 주식회사 일룸은 (대리점명)에게 상기 품목을 아래와 같이 할인한 금액으로\n    공급해 주실 것을 요청합니다.",
        "   ● A유형 품목 : 소비자가의 21%",
        "   ● B유형 품목 : 소비자가의 10%",
    ]:
        p = add_para(doc, txt, size=10, space_before=0, space_after=4)
        p.paragraph_format.left_indent = Cm(0.7)

    doc.add_paragraph()
    add_para(doc, "※ 별첨 : 확정견적서", size=10, space_after=20)

    doc.add_paragraph()
    doc.add_paragraph()

    # 서명
    p_sign = doc.add_paragraph()
    p_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_sign.paragraph_format.space_after = Pt(4)
    r = p_sign.add_run("일룸 롯데영등포점  대표  오승균  (인)")
    set_font(r, size=11)

    # 저장
    path = os.path.join(OUT_DIR, "B2B거래물품할인공급요청서_롯데영등포_수정본.docx")
    doc.save(path)
    print(f"✅ 요청서 저장: {path}")
    return path


if __name__ == "__main__":
    p1 = make_hyeopjojeon()
    p2 = make_b2b_request()
    print("\n📂 생성된 파일:")
    print(f"  {p1}")
    print(f"  {p2}")
