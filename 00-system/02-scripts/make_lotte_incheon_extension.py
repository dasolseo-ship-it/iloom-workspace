# -*- coding: utf-8 -*-
"""
롯데인천점 팝업 계약 연장 품의서 Word 생성
롯데영등포 계약연장 품의 양식 기준
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = r"C:\Users\FURSYS\Downloads"

# ── 헬퍼 ──────────────────────────────────────────
def set_font(run, name="맑은 고딕", size=10, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), name)
    rPr.insert(0, rFonts)

def add_para(doc, text="", align=WD_ALIGN_PARAGRAPH.LEFT, bold=False,
             size=10, space_before=0, space_after=4, indent=0, color=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    if text:
        run = p.add_run(text)
        set_font(run, size=size, bold=bold, color=color)
    return p

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def cell_text(cell, text, size=9, bold=False,
              align=WD_ALIGN_PARAGRAPH.CENTER, color=None, valign=WD_ALIGN_VERTICAL.CENTER):
    cell.vertical_alignment = valign
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)

def border_side(el, side, val="single", sz="6", color="AAAAAA"):
    b = OxmlElement(f"w:{side}")
    b.set(qn("w:val"), val)
    b.set(qn("w:sz"), sz)
    b.set(qn("w:space"), "0")
    b.set(qn("w:color"), color)
    el.append(b)

def set_cell_border_all(cell, color="BBBBBB"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcB = OxmlElement("w:tcBorders")
    for side in ["top","left","bottom","right"]:
        border_side(tcB, side, color=color)
    tcPr.append(tcB)

# ══════════════════════════════════════════════════
#  품의서 생성
# ══════════════════════════════════════════════════
def make_pumui():
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # ── 결재란 + 제목 ──
    # 결재란은 우측 정렬로 단순하게 표현
    top_tbl = doc.add_table(rows=3, cols=7)
    top_tbl.alignment = WD_TABLE_ALIGNMENT.RIGHT
    top_tbl.style = "Table Grid"

    # 헤더행: 작성 / 검토(I II III IV) / 승인
    merge_labels = [
        (0, 0, "작성"),
        (0, 1, "검  토"),   # I~IV 병합
        (0, 5, "승인"),
    ]
    # 1행: 라벨
    r0 = top_tbl.rows[0].cells
    cell_text(r0[0], "작성", size=8, bold=True)
    set_cell_bg(r0[0], "F0F0F0")
    # 검토 I~IV 병합
    r0[1].merge(r0[4])
    cell_text(r0[1], "검   토", size=8, bold=True)
    set_cell_bg(r0[1], "F0F0F0")
    cell_text(r0[5], "승인", size=8, bold=True)
    set_cell_bg(r0[5], "F0F0F0")

    # 2행: 서브 라벨 (I II III IV)
    r1 = top_tbl.rows[1].cells
    cell_text(r1[0], "", size=8)
    for i, lbl in enumerate(["Ⅰ","Ⅱ","Ⅲ","Ⅳ"], 1):
        cell_text(r1[i], lbl, size=8, bold=True)
        set_cell_bg(r1[i], "FAFAFA")
    cell_text(r1[5], "", size=8)

    # 3행: 서명
    r2 = top_tbl.rows[2].cells
    cell_text(r2[0], "서다솔\n05/27", size=8)
    cell_text(r2[1], "안권희\n05/27", size=8)  # 검토I
    cell_text(r2[2], "한병화\n05/27", size=8)  # 검토II
    cell_text(r2[3], "", size=8)
    cell_text(r2[4], "", size=8)
    cell_text(r2[5], "", size=8)  # 승인 (빈칸)

    for row in top_tbl.rows:
        for cell in row.cells:
            set_cell_border_all(cell, "CCCCCC")

    col_widths = [Cm(2.0), Cm(2.0), Cm(2.0), Cm(2.0), Cm(2.0), Cm(2.0)]
    for i, w in enumerate(col_widths):
        for row in top_tbl.rows:
            if i < len(row.cells):
                row.cells[i].width = w

    doc.add_paragraph()

    # ── 문서 제목 ──
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(12)
    r = p_title.add_run("품  의  서  (자유양식)")
    set_font(r, size=20, bold=True)

    # ── 문서 정보 표 ──
    info_tbl = doc.add_table(rows=3, cols=4)
    info_tbl.style = "Table Grid"
    info_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    info_widths = [Cm(2.5), Cm(8.0), Cm(2.5), Cm(4.5)]
    for row in info_tbl.rows:
        for i, w in enumerate(info_widths):
            row.cells[i].width = w
        row.height = Pt(22)

    def info_row(ri, l1, v1, l2, v2, merge_v1=False):
        row = info_tbl.rows[ri]
        cell_text(row.cells[0], l1, bold=True, size=9)
        set_cell_bg(row.cells[0], "F0F0F0")
        if merge_v1:
            row.cells[1].merge(row.cells[2])
        cell_text(row.cells[1], v1, size=9, align=WD_ALIGN_PARAGRAPH.LEFT)
        if not merge_v1:
            cell_text(row.cells[2], l2, bold=True, size=9)
            set_cell_bg(row.cells[2], "F0F0F0")
        cell_text(row.cells[3], v2, size=9, align=WD_ALIGN_PARAGRAPH.LEFT)

    info_row(0, "문서번호", "일룸-품의26-05-     ", "작성일", "2026-05-27")
    info_row(1, "작성부서", "일룸사업부 > 영업개발부문 > 리테일사업팀", "작성자", "서다솔", merge_v1=False)
    info_row(2, "제    목",
             "롯데백화점 인천점 팝업매장 계약 연장의 건(~26.07.31)",
             "열람권한", "부서")

    doc.add_paragraph()

    # ── 인사말 ──
    add_para(doc, "아래와 같이 롯데백화점 인천점 팝업매장 계약 연장을 진행하고자 하오니 검토 후 재가바랍니다.",
             size=10, space_after=8)
    add_para(doc, "----- 아   래 -----",
             align=WD_ALIGN_PARAGRAPH.CENTER, size=10, space_after=12)

    # ── 1. 배경 ──
    add_para(doc, "1. 배경", bold=True, size=11, space_after=4)
    bgs = [
        "1) 일룸 롯데인천점 팝업매장은 2026년 05월 31일까지 계약하여 운영 중",
        "    (관련 품의 : 롯데백화점 인천점 팝업 신규 개설 및 위탁운영 계약 체결의 건)",
        "2) 매장의 계약 만료가 도래하여 추가 계약 연장을 진행하여 상권 매출 활성화를 도모하고자 함",
    ]
    for t in bgs:
        add_para(doc, t, size=10, space_before=0, space_after=3, indent=0.3)

    doc.add_paragraph()

    # ── 2. 내용 ──
    add_para(doc, "2. 내용", bold=True, size=11, space_after=4)
    add_para(doc, "1) 위탁운영 계약 정보", size=10, space_after=4, indent=0.3)

    # 계약 정보 표
    ct = doc.add_table(rows=6, cols=3)
    ct.style = "Table Grid"
    ct.alignment = WD_TABLE_ALIGNMENT.CENTER

    cw = [Cm(3.0), Cm(8.5), Cm(5.0)]
    for row in ct.rows:
        for i, w in enumerate(cw):
            row.cells[i].width = w

    def cinfo(ri, label, content, note="", label_bg="F0F0F0"):
        row = ct.rows[ri]
        cell_text(row.cells[0], label, bold=True, size=9)
        set_cell_bg(row.cells[0], label_bg)
        cell_text(row.cells[1], content, size=9,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_text(row.cells[2], note, size=8,
                  align=WD_ALIGN_PARAGRAPH.LEFT,
                  color=(80,80,80) if note else None)
        for c in row.cells:
            set_cell_border_all(c, "BBBBBB")

    cinfo(0, "매장명",   "일룸 롯데인천점 팝업")
    cinfo(1, "매장유형", "입점 BS(백화점)",    "위탁판매계약 체결")
    cinfo(2, "운영주체", "이종균 대표")
    cinfo(3, "계약기간", "2026.06.01 ~ 2026.07.31 (2개월)",
          "※ 대리점 위탁계약은 일룸 기준으로 체결예정")

    # 판매수수료 행 (복합 내용)
    fee_row = ct.rows[4]
    cell_text(fee_row.cells[0], "판매수수료", bold=True, size=9)
    set_cell_bg(fee_row.cells[0], "F0F0F0")
    set_cell_border_all(fee_row.cells[0], "BBBBBB")

    fee_content = (
        "기본 판매수수료   A유형·B유형   21%\n"
        "추가 판매수수료   입점 수수료 15% 구간   4%\n"
        "총 판매수수료   25%"
    )
    cell_text(fee_row.cells[1], fee_content, size=9)
    set_cell_border_all(fee_row.cells[1], "BBBBBB")
    cell_text(fee_row.cells[2],
              "※ 입점 BS(백화점) 기준과 동일\n(참고: 입점물(백화점) 유통망 대리점 유형 신설의 건)",
              size=8, align=WD_ALIGN_PARAGRAPH.LEFT, color=(80,80,80))
    set_cell_border_all(fee_row.cells[2], "BBBBBB")

    cinfo(5, "거래보증금", "거래보증금 2천만원", "※ 기존 보증금 유지")

    doc.add_paragraph()

    # 2) 계약 진행
    add_para(doc, "2) 계약 진행", size=10, space_after=4, indent=0.3)
    for t in [
        "① 품의 결재 득 후 전자 계약서 체결 예정 → 이종균 대표(모두싸인)",
        "※ 롯데백화점과의 입점 계약은 신유통개발팀에서 별도 품의 후 체결",
    ]:
        add_para(doc, t, size=10, space_before=0, space_after=3, indent=0.7)

    doc.add_paragraph()

    # ── 3. 첨부 ──
    add_para(doc, "3. 첨부", bold=True, size=11, space_after=4)
    add_para(doc, "1) 일룸 대리점 26년 계약정서_롯데인천260601-260731",
             size=10, space_after=3, indent=0.3)

    # 저장
    path = os.path.join(OUT_DIR, "품의서_롯데인천점팝업_계약연장_260731.docx")
    doc.save(path)
    print(f"✅ 저장 완료: {path}")
    return path


if __name__ == "__main__":
    make_pumui()
