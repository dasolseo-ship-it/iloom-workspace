import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT_PATH = r"c:\Users\FURSYS\Downloads\iloom-workspace-claude\50-resources\롯데바이오로직스 공동구매 프로모션 합의서.docx"

stores = [
    {
        "name": "송도5",
        "type": "투자형B · 6:4",
        "ratio_text": "각 대리점과 본사가 각각 40%(대리점), 60%(본사)의 비율로 분담한다.",
        "dealer_company": "(주)알루컴 인천점",
        "dealer_biz_no": "105-85-41701",
        "dealer_addr": "인천 남동구 인주대로 644",
        "dealer_rep": "이 종 균",
    },
    {
        "name": "인천중앙2",
        "type": "일반B · 5:5",
        "ratio_text": "각 대리점과 본사가 각각 50%의 비율로 분담한다.",
        "dealer_company": "(주)알루컴 인천점",
        "dealer_biz_no": "105-85-41701",
        "dealer_addr": "인천 남동구 인주대로 644",
        "dealer_rep": "이 종 균",
    },
    {
        "name": "인천검단",
        "type": "투자형B · 6:4",
        "ratio_text": "각 대리점과 본사가 각각 40%(대리점), 60%(본사)의 비율로 분담한다.",
        "dealer_company": "일룸 인천 검단점",
        "dealer_biz_no": "872-03-02670",
        "dealer_addr": "인천 서구 원당대로 847, 4층",
        "dealer_rep": "유 종 규",
    },
    {
        "name": "신세계시흥2",
        "type": "투자형B · 6:4",
        "ratio_text": "각 대리점과 본사가 각각 40%(대리점), 60%(본사)의 비율로 분담한다.",
        "dealer_company": "일룸 신세계시흥",
        "dealer_biz_no": "133-40-01343",
        "dealer_addr": "경기도 시흥시 서해안로 699, 2층 (배곧동, 시흥 프리미엄 아울렛)",
        "dealer_rep": "최 하 은",
    },
    {
        "name": "김포5",
        "type": "일반B · 5:5",
        "ratio_text": "각 대리점과 본사가 각각 50%의 비율로 분담한다.",
        "dealer_company": "일룸김포점",
        "dealer_biz_no": "137-17-39033",
        "dealer_addr": "경기 김포시 풍무동 27-4 1,2층",
        "dealer_rep": "장 미 영",
    },
    {
        "name": "부천3",
        "type": "일반B · 5:5",
        "ratio_text": "각 대리점과 본사가 각각 50%의 비율로 분담한다.",
        "dealer_company": "일룸 부천점",
        "dealer_biz_no": "313-36-00088",
        "dealer_addr": "경기 부천시 길주로 462",
        "dealer_rep": "최 민 영",
    },
]

def set_font(run, size=11, bold=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

def add_paragraph(doc, text="", bold=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=4):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        set_font(run, size=size, bold=bold)
    return p

def add_signature_row(doc, label, company, biz_no, addr, rep):
    for line, content in [
        (f'"{label}"', f"상호 : {company} ( {biz_no} )"),
        ("", f"주소 : {addr}"),
        ("", f"대표 : {rep}                                          (인)"),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(f"{label_pad(label, line):<8}")
        set_font(r1, bold=bool(line))
        r2 = p.add_run(content)
        set_font(r2)

def label_pad(original, current):
    return current if current else ""

def build_agreement(doc, store, is_first):
    if not is_first:
        doc.add_page_break()

    # 제목
    add_paragraph(doc, "임직원 할인판매 프로모션 합의서", bold=True, size=16,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=6)

    # 매장 정보 부제
    add_paragraph(doc, f"[ {store['name']}  |  {store['type']} ]",
                  bold=False, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

    # 머리글
    header = (
        f"본 합의서는 ㈜일룸 (이하 공급업자)과 {store['name']} (이하 대리점) 간에\n"
        "다음과 같이 롯데바이오로직스 임직원 할인판매 프로모션 관련 진행에 대하여 합의한다."
    )
    add_paragraph(doc, header, size=11, space_after=4)
    add_paragraph(doc, "- 다         음 -", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

    # 조항 내용
    articles = [
        ("제1조 [목적]", [
            "1) 이 합의는 롯데바이오로직스 임직원 할인판매 프로모션기간 중 분담금 협의 조건 등을 정함을 목적으로 한다.",
        ]),
        ("제2조 [수수료 변경 품목]", [
            "1) 공급업자가 정한 일룸 전체 제품을 수수료 변경 품목으로 한다.",
            "2) 공급품목은 운영 중 상호합의하에 조정할 수 있으며, 계약기간 중 공급품목의 변동은 없으나 불가피한 사유\n   발생 시 사전 통지 후 공급을 중단할 수 있다.",
        ]),
        ("제3조 [판매수수료 및 미수금]", [
            "1) 본 프로모션에 대한 판매수수료율은 「2025년 위탁판매대리점 판매수수료 및 지원금 약정서」에서 정한 판매수수료율을\n   동일하게 적용한다. 다만, 판매수수료 산정의 기준이 되는 금액은 정상 소비자가(VAT별도)에서 10%를 할인한 금액으로 한다.",
            f"2) 제1항의 10% 할인 금액(VAT 별도)에 대하여는 {store['ratio_text']}",
            "3) 전항에 따라 발생하는 대리점 부담분은 미수금으로 산정하여 위탁판매수수료(VAT 별도)정산 시 상계 처리하기로 한다.\n   상기 정산 방식에 대해 설명받고 이에 동의하였음을 상호간에 확인한다.",
            "4) 대리점은 본 조의 내용에 대하여 충분한 설명을 제공받았으며,\n   이를 명확히 이해한 상태에서 전적으로 자유로운 의사에 따라 본 조항에 동의하였음을 당사자 상호 간에 확인한다.",
        ]),
        ("제4조 [프로모션 기간 및 상세내용]", [
            "1) 롯데바이오로직스 임직원 할인판매 프로모션은 일정 기간만 한정 진행되는 건을 대상으로 한다.\n   임직원 할인판매 프로모션 적용 기간\n   (수주 기준: 2026년 06월 13일 ~ 06월 28일 / 매출 기준: 2026년 06월 13일 ~ 08월 31일)",
            "2) 롯데바이오로직스 임직원 사원증 및 명함 확인은 필수이며, 구매횟수는 1인 1건으로 제한합니다.",
        ]),
        ("제5조 [분쟁해결 및 재판관할]", [
            "1) 이 합의에서 정하지 아니한 사항 또는 이 합의의 내용에 대하여 공급업자와 대리점간 분쟁이 발생한 경우,\n   공급업자와 대리점의 합의된 의사에 따른다. 공급업자와 대리점사이에 내용 해석과 관련한 합의가 이루어지지\n   않은 경우 일반 상관례 및 관련 법령에 따른다.",
            "2) 제1항에 의해서도 이 합의와 관련한 분쟁이 해결되지 아니한 경우에는 「대리점거래의 공정화에 관한 법률」\n   제19조 제1항에 따라 대리점분쟁조정협의회에 조정을 신청할 수 있다. 이 경우에 공급업자와 대리점은 조정절차에\n   성실하게 임하며, 원활한 분쟁해결을 위해 노력한다.",
            "3) 공급업자와 대리점이 제2항에 따른 조정을 신청하지 아니하는 경우, 이 합의에 관한 분쟁의 해결은 통상의\n   민사재판에 의하며, 관할은 민사소송법에 따른다.",
        ]),
        ("제6조 [합의의 효력 등]", [
            "1) 공급업자와 대리점은 이 합의를 체결하기 전에 충분한 협의를 거쳤고, 합의 내용을 모두 숙지하였으며,\n   이 합의를 증명하기 위하여 전자계약으로 서명 또는 기명날인하여 각각 1부씩 보관한다.",
            "2) 이 합의서에 기재된 내용만이 공급업자와 대리점 사이에 합의된 내용이며, 이 이외의 내용에 대한 당사자 간의\n   그 어떠한 구두 합의도 당사자를 구속하지 아니한다.",
            "3) 이 합의서의 내용은 공급업자와 대리점 사이의 서면 합의에 의해서만 변경되거나 수정될 수 있으며,\n   그 변경 및 수정은 공급업자와 대리점이 해당 서면에 서명 또는 기명날인함과 동시에 그 효력을 발생한다.",
        ]),
    ]

    for title, items in articles:
        add_paragraph(doc, title, bold=True, size=11, space_before=6, space_after=2)
        for item in items:
            add_paragraph(doc, item, size=11, space_before=0, space_after=2)

    # 날짜
    add_paragraph(doc, "2026년        월        일", align=WD_ALIGN_PARAGRAPH.CENTER,
                  space_before=14, space_after=10)

    # 서명란
    for label, company, biz_no, addr, rep in [
        ("공급업자", "주식회사 일룸", "215-86-93600",
         "서울특별시 송파구 오금동 45-1 퍼시스빌딩 3층", "정 보 은"),
        ("대리점", store["dealer_company"], store["dealer_biz_no"],
         store["dealer_addr"], store["dealer_rep"]),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(f'"{label}"')
        set_font(r, bold=True)

        for line in [
            f"상호 : {company} ( {biz_no} )",
            f"주소 : {addr}",
            f"대표 : {rep}                                    (인)",
        ]:
            p2 = doc.add_paragraph()
            p2.paragraph_format.space_before = Pt(0)
            p2.paragraph_format.space_after = Pt(1)
            p2.paragraph_format.left_indent = Cm(1.5)
            r2 = p2.add_run(line)
            set_font(r2)

        add_paragraph(doc, "", space_after=6)


doc = Document()

# 여백 설정
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)

for i, store in enumerate(stores):
    build_agreement(doc, store, is_first=(i == 0))

doc.save(OUTPUT_PATH)
print(f"저장 완료: {OUTPUT_PATH}")
