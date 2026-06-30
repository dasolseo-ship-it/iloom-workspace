import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# 페이지 여백 설정
section = doc.sections[0]
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# 기본 스타일 설정
style = doc.styles['Normal']
style.font.name = 'Malgun Gothic'
style.font.size = Pt(9)
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

def set_cell_bg(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def set_cell_font(cell, text, bold=False, size=9, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ''
    para = cell.paragraphs[0]
    para.alignment = align
    run = para.add_run(text)
    run.font.name = 'Malgun Gothic'
    run.font.bold = bold
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    if color:
        run.font.color.rgb = RGBColor(*color)

def set_col_widths(table, widths_cm):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                cell.width = Cm(widths_cm[i])

def add_border_to_table(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'AAAAAA')
        tblBorders.append(border)
    tblPr.append(tblBorders)

# ───────────────────────────────────────────────
# 제목
# ───────────────────────────────────────────────
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_para.add_run('품의서 (자유양식)')
title_run.font.name = 'Malgun Gothic'
title_run.font.bold = True
title_run.font.size = Pt(18)
title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

doc.add_paragraph()

# ───────────────────────────────────────────────
# 결재 테이블 (우측 정렬)
# ───────────────────────────────────────────────
approval_table = doc.add_table(rows=2, cols=6)
approval_table.alignment = WD_TABLE_ALIGNMENT.RIGHT
add_border_to_table(approval_table)
headers = ['작성', '검토 I', '검토 II', '검토 III', '검토 IV', '승인']
for i, h in enumerate(headers):
    set_cell_font(approval_table.rows[0].cells[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(approval_table.rows[0].cells[i], 'F0F0F0')
for i in range(6):
    approval_table.rows[1].cells[i].text = ''
    approval_table.rows[1].cells[i].add_paragraph()
    approval_table.rows[1].cells[i].add_paragraph()

doc.add_paragraph()

# ───────────────────────────────────────────────
# 문서 정보 테이블
# ───────────────────────────────────────────────
info_table = doc.add_table(rows=3, cols=4)
add_border_to_table(info_table)

rows_data = [
    ('문서번호', '일룸-품의26-XX-XXXXX', '작성일', '2026-06-30'),
    ('작성부서', '일룸사업부 > 영업개발부문 > 대리점파트', '작성자', '서다솔'),
    ('제목', '전주 서신더샵비발디 입주 공략의 건(중화산4점)(입주박람회)', '열람권한', '부서'),
]
for r_idx, (l1, v1, l2, v2) in enumerate(rows_data):
    row = info_table.rows[r_idx]
    set_cell_font(row.cells[0], l1, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(row.cells[0], 'F0F0F0')
    set_cell_font(row.cells[1], v1)
    set_cell_font(row.cells[2], l2, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(row.cells[2], 'F0F0F0')
    set_cell_font(row.cells[3], v2)

set_col_widths(info_table, [2.5, 8.0, 2.0, 3.5])

doc.add_paragraph()

intro = doc.add_paragraph('아래와 같이 전주 서신더샵비발디 입주아파트를 공략하고자 하오니 검토 후 재가 바랍니다.')
intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
intro.runs[0].font.name = 'Malgun Gothic'
intro.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

sep = doc.add_paragraph('───── 아   래 ─────')
sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
sep.runs[0].font.name = 'Malgun Gothic'
sep.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

# ───────────────────────────────────────────────
# 섹션 제목 헬퍼
# ───────────────────────────────────────────────
def add_section_title(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Malgun Gothic'
    run.font.bold = True
    run.font.size = Pt(10)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    run.font.color.rgb = RGBColor(0xC8, 0x0A, 0x1E)
    return p

def add_note(doc, text):
    p = doc.add_paragraph(text)
    p.runs[0].font.name = 'Malgun Gothic'
    p.runs[0].font.size = Pt(8)
    p.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    p.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    return p

# ───────────────────────────────────────────────
# ※ 본 입주공략 요약
# ───────────────────────────────────────────────
add_section_title(doc, '※ 본 입주공략 관련 요약')

summary_table = doc.add_table(rows=6, cols=3)
add_border_to_table(summary_table)

summary_data = [
    ('구분', '내용', '비고'),
    ('공략대상지', '전주 서신더샵비발디', '중화산4점 단독 공략'),
    ('입주 세대 수 및 시점', '[세대수 확인 필요] / 2026년 10월 입주', ''),
    ('공략지 목표', '입주박람회 매출 확대 및 고객 접점 강화', 'VAT별도'),
    ('운영 프로모션', '추가혜택(박람회 전용): 세트구매 조건 충족 시 구매금액 5% 추가', '본사:대리점 = 5:5 부담'),
    ('비고', '대리점 단독 응대 (전시품 별도 이동 없음)', ''),
]
for r_idx, row_data in enumerate(summary_data):
    row = summary_table.rows[r_idx]
    for c_idx, text in enumerate(row_data):
        bold = (r_idx == 0)
        align = WD_ALIGN_PARAGRAPH.CENTER if (r_idx == 0 or c_idx == 0) else WD_ALIGN_PARAGRAPH.LEFT
        set_cell_font(row.cells[c_idx], text, bold=bold, align=align)
        if r_idx == 0 or c_idx == 0:
            set_cell_bg(row.cells[c_idx], 'F0F0F0')

set_col_widths(summary_table, [3.5, 9.5, 3.0])
doc.add_paragraph()

# ───────────────────────────────────────────────
# 1. 배경 및 목적
# ───────────────────────────────────────────────
add_section_title(doc, '1. 배경 및 목적')
for item in [
    '1) 전주 서신더샵비발디(2026년 10월 입주 예정) 입주 수요 공략을 통해 일룸 중화산4점 매출 확대 도모',
    '2) 입주박람회 참가를 통해 라이프스타일 가구 브랜드 일룸 인지도 제고 및 지역 고객 접점 강화',
]:
    p = doc.add_paragraph(item)
    p.runs[0].font.name = 'Malgun Gothic'
    p.runs[0].font.size = Pt(9)
    p.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
doc.add_paragraph()

# ───────────────────────────────────────────────
# 2. 공략대상
# ───────────────────────────────────────────────
add_section_title(doc, '2. 공략대상')

target_table = doc.add_table(rows=6, cols=3)
add_border_to_table(target_table)
target_data = [
    ('구분', '내용', '비고'),
    ('아파트명', '전주 서신더샵비발디', ''),
    ('위치', '전주시 덕진구 서신동 일원', ''),
    ('입주 월', '2026년 10월', ''),
    ('세대수', '[확인 필요]', ''),
    ('공략 매장', '일룸 중화산4점', ''),
]
for r_idx, row_data in enumerate(target_data):
    row = target_table.rows[r_idx]
    for c_idx, text in enumerate(row_data):
        bold = (r_idx == 0)
        align = WD_ALIGN_PARAGRAPH.CENTER if (r_idx == 0 or c_idx == 0) else WD_ALIGN_PARAGRAPH.LEFT
        set_cell_font(row.cells[c_idx], text, bold=bold, align=align)
        if r_idx == 0 or c_idx == 0:
            set_cell_bg(row.cells[c_idx], 'F0F0F0')
set_col_widths(target_table, [3.5, 9.5, 3.0])

doc.add_paragraph()
p = doc.add_paragraph('공략대상지 정보')
p.runs[0].font.bold = True
p.runs[0].font.name = 'Malgun Gothic'
p.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

for item in [
    '- 박람회 부스 규모: 2.5m × 2.5m',
    '- 입주박람회 일정: 2026년 8월 중 (구체 일정 확인 필요)',
    '- 운영 방식: 현장 사장님 단독 응대, 전시품 별도 이동 없음',
]:
    p = doc.add_paragraph(item)
    p.runs[0].font.name = 'Malgun Gothic'
    p.runs[0].font.size = Pt(9)
    p.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
doc.add_paragraph()

# ───────────────────────────────────────────────
# 3. 홍보방안
# ───────────────────────────────────────────────
add_section_title(doc, '3. 홍보방안')
p = doc.add_paragraph('1) 입주박람회')
p.runs[0].font.bold = True
p.runs[0].font.name = 'Malgun Gothic'
p.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
add_note(doc, '(단위: 원, VAT별도)')

cost_table = doc.add_table(rows=6, cols=7)
add_border_to_table(cost_table)
cost_headers = ['구분', '총 비용', '본사', '분담률', '대리점', '분담률', '비고']
cost_data = [
    ('참가비(입점료)', '3,500,000', '1,750,000', '50%', '1,750,000', '50%', ''),
    ('입주리플릿(3단 접지)', '300,000', '150,000', '50%', '150,000', '50%', '디곳통 디자인·인성 인쇄, 예가'),
    ('롤업 배너 시안 제작', '100,000', '50,000', '50%', '50,000', '50%', '기존 배너 활용, 시안만 교체'),
    ('박람회 경품', '200,000', '100,000', '50%', '100,000', '50%', '부가세 없음, 예가'),
    ('합계 (프로모션 제외)', '4,100,000', '2,050,000', '50%', '2,050,000', '50%', ''),
]
for i, h in enumerate(cost_headers):
    set_cell_font(cost_table.rows[0].cells[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(cost_table.rows[0].cells[i], 'F0F0F0')
for r_idx, row_data in enumerate(cost_data):
    row = cost_table.rows[r_idx + 1]
    for c_idx, text in enumerate(row_data):
        is_total = (r_idx == 4)
        align = WD_ALIGN_PARAGRAPH.CENTER if c_idx in [1,2,3,4,5] else WD_ALIGN_PARAGRAPH.LEFT
        set_cell_font(row.cells[c_idx], text, bold=is_total, align=align)
        if is_total:
            set_cell_bg(row.cells[c_idx], 'F5F5F5')

set_col_widths(cost_table, [3.5, 2.2, 2.2, 1.5, 2.2, 1.5, 3.0])

add_note(doc, '※ 프로모션비용(추가혜택 5%)은 예상 매출 확정 후 별도 채산 반영 예정')
add_note(doc, '※ 박람회 운전장비 및 퀵 비용은 지원하지 아니함')
doc.add_paragraph()

# ───────────────────────────────────────────────
# 4. 프로모션
# ───────────────────────────────────────────────
add_section_title(doc, '4. 프로모션')
p = doc.add_paragraph('1) 입주박람회 현장 프로모션 — 추가혜택')
p.runs[0].font.bold = True
p.runs[0].font.name = 'Malgun Gothic'
p.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
add_note(doc, '(금액단위: VAT포함, 혜택 금액 부가세 없음)')

promo_intro = doc.add_paragraph('아래 조건 충족 시 구매금액의 5% 추가 혜택 적용  |  본사 : 대리점 = 5 : 5 부담')
promo_intro.runs[0].font.name = 'Malgun Gothic'
promo_intro.runs[0].font.bold = True
promo_intro.runs[0].font.size = Pt(9)
promo_intro.runs[0].font.color.rgb = RGBColor(0xC8, 0x0A, 0x1E)
promo_intro.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

promo_table = doc.add_table(rows=4, cols=2)
add_border_to_table(promo_table)
promo_data = [
    ('공간', '대상 품목'),
    ('다이닝', '토스카노 + 휴스턴 세트 구매'),
    ('옷장', '옷장 몸통 3통 이상 구매'),
    ('소파', '3인 이상 (1인 블케 포함)'),
]
for r_idx, (l, v) in enumerate(promo_data):
    row = promo_table.rows[r_idx]
    bold = (r_idx == 0)
    align_l = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_font(row.cells[0], l, bold=bold, align=align_l)
    set_cell_font(row.cells[1], v, bold=bold, align=WD_ALIGN_PARAGRAPH.LEFT)
    if r_idx == 0:
        set_cell_bg(row.cells[0], 'F0F0F0')
        set_cell_bg(row.cells[1], 'F0F0F0')
set_col_widths(promo_table, [3.5, 12.5])

add_note(doc, '※ 200만원 구간 추가혜택 적용 여부는 채산 검토 후 대리점과 협의 예정')
doc.add_paragraph()

p = doc.add_paragraph('2) 공통사항')
p.runs[0].font.bold = True
p.runs[0].font.name = 'Malgun Gothic'
p.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

for item in [
    '- 납기한도: 2026년 10월 31일 (입주 후 1개월)',
    '- 고객 혜택금은 천원단위 절사 지급 기준',
    '- 쿠시노반짝(및 이후 브랜드 스팟성 프로모션 포함) / 프레임반값과 중복 불가',
    '- 납기 종료 후 익월말 포인트 일괄 지급 (영업담당자 별도 정산)',
]:
    p = doc.add_paragraph(item)
    p.runs[0].font.name = 'Malgun Gothic'
    p.runs[0].font.size = Pt(9)
    p.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
doc.add_paragraph()

# ───────────────────────────────────────────────
# 5. 예상매출 및 손익
# ───────────────────────────────────────────────
add_section_title(doc, '5. 예상매출 및 손익')
add_note(doc, '(단위: 원, VAT별도)')

profit_table = doc.add_table(rows=4, cols=3)
add_border_to_table(profit_table)
profit_data = [
    ('구분', '금액', '비고'),
    ('예상매출', '[세대수 확정 후 기입]', '공략기간 목표 매출'),
    ('공략비용 (고정)', '4,100,000', '입점료+리플릿+배너시안+경품'),
    ('채산이익', '[매출 확정 후 기입]', ''),
]
for r_idx, row_data in enumerate(profit_data):
    row = profit_table.rows[r_idx]
    for c_idx, text in enumerate(row_data):
        bold = (r_idx == 0)
        align = WD_ALIGN_PARAGRAPH.CENTER if (r_idx == 0 or c_idx == 0) else WD_ALIGN_PARAGRAPH.LEFT
        set_cell_font(row.cells[c_idx], text, bold=bold, align=align)
        if r_idx == 0 or c_idx == 0:
            set_cell_bg(row.cells[c_idx], 'F0F0F0')
set_col_widths(profit_table, [3.5, 5.0, 7.5])
doc.add_paragraph()

# ───────────────────────────────────────────────
# 6. 기타사항
# ───────────────────────────────────────────────
add_section_title(doc, '6. 기타사항')
for item in [
    '1) 특정 기간에만 적용 가능한 프로모션으로 "일룸 중화산4점"에서 진행',
    '2) 아파트 약명: 서신더샵비발디 (건명 기입 기준)',
    '3) 기존 보관 중인 롤업 배너(전군 보관) 활용 — 시안만 제작하여 중화산4점으로 발송',
    '4) 3단 접지리플릿은 디곳통 디자인 후 인성 인쇄 발주 예정',
]:
    p = doc.add_paragraph(item)
    p.runs[0].font.name = 'Malgun Gothic'
    p.runs[0].font.size = Pt(9)
    p.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
doc.add_paragraph()

# ───────────────────────────────────────────────
# 7. 첨부파일
# ───────────────────────────────────────────────
add_section_title(doc, '7. 첨부파일')
p = doc.add_paragraph('- 없음 (추후 입주박람회 합의서 체결 내용에 따라 전자계약 진행 예정)')
p.runs[0].font.name = 'Malgun Gothic'
p.runs[0].font.size = Pt(9)
p.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

doc.add_paragraph()
p = doc.add_paragraph('끝.')
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.runs[0].font.bold = True
p.runs[0].font.name = 'Malgun Gothic'
p.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

# 저장
out_path = r'20-operations/24-입주공략/전주서신더샵비발디/2026-08_전주서신더샵비발디_입주공략_품의_초안.docx'
doc.save(out_path)
print(f'저장 완료: {out_path}')
