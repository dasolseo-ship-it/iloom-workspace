# -*- coding: utf-8 -*-
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# 기본 여백 설정
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

def set_font(run, size=10, bold=False, color=None):
    run.font.name = 'Noto Sans KR'
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    rPr.insert(0, rFonts)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    if level == 1:
        set_font(run, size=15, bold=True, color=(0, 0, 0))
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        set_font(run, size=11, bold=True, color=(0, 0, 0))
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
    return p

def add_body(doc, text, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    set_font(run, size=10)
    p.paragraph_format.space_after = Pt(2)
    return p

def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    set_font(run, size=9, color=(100, 100, 100))
    p.paragraph_format.space_after = Pt(4)
    return p

def set_cell(cell, text, bold=False, center=False, bg_color=None, font_color=None, size=9):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.clear()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=font_color)
    if bg_color:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), bg_color)
        tcPr.append(shd)

# ===== 제목 =====
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run('LXH목동점 위탁판매 계약 종료\n채권 정산보고서')
set_font(title_run, size=18, bold=True)
title_p.paragraph_format.space_before = Pt(6)
title_p.paragraph_format.space_after = Pt(4)

# 구분선
doc.add_paragraph('─' * 50)

# 메타 정보
meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
meta_run = meta_p.add_run('작성일: 2025년 9월 30일    작성자: 서다솔 (대리점파트)    문서구분: 채권대사 정산보고')
set_font(meta_run, size=9, color=(100, 100, 100))

# ===== 1. 목적 =====
add_heading(doc, '1. 목적', level=2)
add_body(doc, '일룸 LXH목동점(투자형 B-Shop)의 위탁판매 대리점 계약이 2025년 9월 30일부로 종료됨에 따라,')
add_body(doc, '위탁판매 기간(2019년 1월 ~ 2024년 11월) 중 발생한 미수금을 수주 대조하여 최종 채권 잔액을 정산·보고함.')

# ===== 2. 대리점 기본정보 =====
add_heading(doc, '2. 대리점 기본정보', level=2)

t = doc.add_table(rows=6, cols=2)
t.style = 'Table Grid'
t.autofit = False
t.columns[0].width = Cm(4.5)
t.columns[1].width = Cm(10.5)

headers2 = [('대리점명', 'LXH목동점 (투자형 B-Shop)'),
            ('대리점코드', '61LM96'),
            ('주소', '서울 양천구 신월로 323'),
            ('대표자', '구한정'),
            ('사업자등록번호', '493-27-00631'),
            ('계약 종료일', '2025년 9월 30일')]
for i, (k, v) in enumerate(headers2):
    set_cell(t.rows[i].cells[0], k, bold=True, center=True, bg_color='F2F2F2', size=9)
    set_cell(t.rows[i].cells[1], v, size=9)

# ===== 3. 미수금 수주 대조 결과 =====
add_heading(doc, '3. 미수금 수주 대조 결과', level=2)
add_body(doc, '대조 기간: 2019년 1월 ~ 2024년 11월  (단위: 원, VAT 포함)')

t2 = doc.add_table(rows=2, cols=5)
t2.style = 'Table Grid'
t2.autofit = False
col_widths = [3.2, 3.2, 2.8, 2.8, 3.0]
for i, w in enumerate(col_widths):
    for cell in t2.columns[i].cells:
        cell.width = Cm(w)

headers3 = ['a. 외상매출금', 'b. 외상매출금 입금', 'c. 할부 수수료', 'd. 반품 수수료', '진여 채권 (a-b+c+d)']
for i, h in enumerate(headers3):
    set_cell(t2.rows[0].cells[i], h, bold=True, center=True, bg_color='333333', font_color=(255,255,255), size=9)

values = ['4,585,635,931', '4,488,287,436', '88,891,907', '100,900', '△1,442,512']
for i, v in enumerate(values):
    if i == 4:
        set_cell(t2.rows[1].cells[i], v, bold=True, center=True, bg_color='FFE0E0', font_color=(200,0,0), size=9)
    else:
        set_cell(t2.rows[1].cells[i], v, center=True, size=9)

add_note(doc, '※ 진여채권 산정 상세 (ERP 기준)')
add_note(doc, '   • ERP 대비 소과(貸越) 처리')
add_note(doc, '   • 총 진여채권 1,442,512원 중 미환급분 980,003원 차충 후 순 진여채권 462,509원')
add_note(doc, '   • 462,509원은 반품 금액 오류, 한붓 누락 등으로 인한 차이 (세부 조정 내역 별첨 참조)')

# ===== 4. 거래처 보조원장 대조 =====
add_heading(doc, '4. 거래처 보조원장 대조 확인', level=2)

add_body(doc, '① 신도림테크노마트점 (41LMS1)  —  기준일: 2025-03-01', indent=True)

t3 = doc.add_table(rows=4, cols=5)
t3.style = 'Table Grid'
t3.autofit = False
h3_cols = ['계정코드', '거래처', '잔금코드', '차변', '잔액']
h3_w = [2.5, 4.5, 2.5, 3.0, 3.0]
for i, w in enumerate(h3_w):
    for cell in t3.columns[i].cells:
        cell.width = Cm(w)
for i, h in enumerate(h3_cols):
    set_cell(t3.rows[0].cells[i], h, bold=True, center=True, bg_color='404040', font_color=(255,255,255), size=9)

t3_data = [
    ('41LMS1', '외상(신도림테크노마트)', '1117000', '115,779,118', '115,779,118'),
    ('41LMS1', '외상매출채권', '2190300', '591,687,363', '591,687,363'),
    ('합계', '', '', '', '707,467,481'),
]
for r, row in enumerate(t3_data, 1):
    for c, val in enumerate(row):
        bold = (r == 3)
        bg = 'F5F5F5' if r == 3 else None
        set_cell(t3.rows[r].cells[c], val, bold=bold, center=(c != 1), bg_color=bg, size=9)

doc.add_paragraph()
add_body(doc, '② 신도림3 (61LM06)  —  기준일: 2025-07-30', indent=True)

t4 = doc.add_table(rows=4, cols=5)
t4.style = 'Table Grid'
t4.autofit = False
for i, w in enumerate(h3_w):
    for cell in t4.columns[i].cells:
        cell.width = Cm(w)
for i, h in enumerate(h3_cols):
    set_cell(t4.rows[0].cells[i], h, bold=True, center=True, bg_color='404040', font_color=(255,255,255), size=9)

t4_data = [
    ('61LM06', '외상매출금', '1112010', '1,442,512', '1,442,512'),
    ('61LM06', '(기타 항목 합산)', '', '499,647,845', '498,249,536'),
    ('합계', '', '', '', '499,691,048'),
]
for r, row in enumerate(t4_data, 1):
    for c, val in enumerate(row):
        bold = (r == 3)
        bg = 'F5F5F5' if r == 3 else None
        set_cell(t4.rows[r].cells[c], val, bold=bold, center=(c != 1), bg_color=bg, size=9)

doc.add_paragraph()
add_body(doc, '③ 대조 결과 요약', indent=True)

t5 = doc.add_table(rows=3, cols=2)
t5.style = 'Table Grid'
t5.columns[0].width = Cm(7)
t5.columns[1].width = Cm(8)
set_cell(t5.rows[0].cells[0], '거래처', bold=True, center=True, bg_color='404040', font_color=(255,255,255), size=9)
set_cell(t5.rows[0].cells[1], '외상매출금 잔액', bold=True, center=True, bg_color='404040', font_color=(255,255,255), size=9)
set_cell(t5.rows[1].cells[0], '신도림3 (61LM06)', center=True, size=9)
set_cell(t5.rows[1].cells[1], '△1,442,512원', bold=True, center=True, font_color=(200, 0, 0), size=9)
set_cell(t5.rows[2].cells[0], '신도림테크노마트 (41LMS1)', center=True, size=9)
set_cell(t5.rows[2].cells[1], '0원', center=True, size=9)

# ===== 5. 채권 정산 처리 내역 =====
add_heading(doc, '5. 채권 정산 처리 내역', level=2)

t6 = doc.add_table(rows=2, cols=5)
t6.style = 'Table Grid'
t6.autofit = False
h6_w = [2.5, 3.0, 2.5, 3.5, 3.5]
for i, w in enumerate(h6_w):
    for cell in t6.columns[i].cells:
        cell.width = Cm(w)
h6 = ['구분', '계정과목', '금액 (원)', '처리 방법', '대상 거래처']
for i, h in enumerate(h6):
    set_cell(t6.rows[0].cells[i], h, bold=True, center=True, bg_color='C80A1E', font_color=(255,255,255), size=9)

t6_vals = ['진여채권 정리', '외상매출금1', '1,442,512', '참이익 처리', '신도림3 (61LM06)']
for i, v in enumerate(t6_vals):
    set_cell(t6.rows[1].cells[i], v, center=True, size=9)

add_note(doc, '※ 액면금: 없음')

# ===== 6. 요청사항 =====
add_heading(doc, '6. 요청사항', level=2)
add_body(doc, '1) 신도림3 (61LM06) 사용승지 처리 → 일룸 재경팀 요청')
add_note(doc, '   • 계약 종료 매장으로 추가 수주 발생 없음')
add_note(doc, '   • 잔여 채권 정리 완료 후 계정 폐쇄 요청')

# ===== 7. 첨부자료 =====
add_heading(doc, '7. 첨부자료', level=2)
add_body(doc, '1) (채권대사) LXH목동점 V.2 — 수주 대조 상세 내역')
add_body(doc, '2) 거래처별 보조원장 (신도림3, 신도림테크노마트)')
add_body(doc, '3) 진여채권 조정 세부내역 (반품 오류·누락 항목)')

# 저장
out_path = r'c:\Users\FURSYS\Downloads\iloom-workspace-claude\20-operations\23-채산관리\2025_LXH목동점_채권정산보고서.docx'
doc.save(out_path)
print(f'saved: {out_path}')
