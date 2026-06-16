# -*- coding: utf-8 -*-
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

def set_font(run, size=10, bold=False, color=None):
    run.font.name = 'Malgun Gothic'
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    rPr.insert(0, rFonts)

def add_heading(doc, text, level=2):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if level == 1:
        set_font(run, size=16, bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
    else:
        set_font(run, size=11, bold=True)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
    return p

def add_body(doc, text, indent=False, size=10):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    set_font(run, size=size)
    p.paragraph_format.space_after = Pt(2)
    return p

def add_note(doc, text, size=9):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    run = p.add_run(text)
    set_font(run, size=size, color=(90, 90, 90))
    p.paragraph_format.space_after = Pt(2)
    return p

def cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell(cell, text, bold=False, center=False, bg=None, fc=None, size=9):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.clear()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=fc)
    if bg:
        cell_bg(cell, bg)

def make_table(doc, headers, rows, col_widths, hdr_bg='333333', hdr_fc=(255,255,255)):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.autofit = False
    for i, w in enumerate(col_widths):
        for c in t.columns[i].cells:
            c.width = Cm(w)
    for i, h in enumerate(headers):
        set_cell(t.rows[0].cells[i], h, bold=True, center=True, bg=hdr_bg, fc=hdr_fc, size=9)
    for ri, row in enumerate(rows, 1):
        for ci, (val, opts) in enumerate(row):
            set_cell(t.rows[ri].cells[ci], val,
                     bold=opts.get('bold', False),
                     center=opts.get('center', False),
                     bg=opts.get('bg'),
                     fc=opts.get('fc'),
                     size=opts.get('size', 9))
    return t

# ===== 제목 =====
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_p.add_run('LXH목동점 위탁판매 계약 종료\n채권 정산보고서')
set_font(r, size=18, bold=True)
title_p.paragraph_format.space_before = Pt(4)
title_p.paragraph_format.space_after = Pt(4)

doc.add_paragraph('─' * 52)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
mr = meta.add_run('작성일: 2025년 9월 30일    작성자: 서다솔 (대리점파트)    문서구분: 채권대사 정산보고')
set_font(mr, size=9, color=(100,100,100))

# ===== 1. 목적 =====
add_heading(doc, '1. 목적')
add_body(doc, '일룸 LXH목동점(투자형 B-Shop)의 위탁판매 대리점 계약이 2025년 9월 30일부로 종료됨에 따라,')
add_body(doc, '위탁판매 계약 기간(2021년 4월 ~ 2025년 9월) 중 발생한 외상매출금 전체를 수주 대조하여')
add_body(doc, '최종 채권 잔액을 정산하고 처리 방침을 보고함.')

# ===== 2. 대리점 기본정보 =====
add_heading(doc, '2. 대리점 기본정보')
info_rows = [
    [('대리점명', {'bold':True,'center':True,'bg':'F2F2F2'}), ('일룸 목동점 (LXH목동) — 투자형 B-Shop', {})],
    [('대리점코드', {'bold':True,'center':True,'bg':'F2F2F2'}), ('61LM96', {})],
    [('주소', {'bold':True,'center':True,'bg':'F2F2F2'}), ('서울 양천구 신월로 323', {})],
    [('대표자', {'bold':True,'center':True,'bg':'F2F2F2'}), ('구한정', {})],
    [('사업자등록번호', {'bold':True,'center':True,'bg':'F2F2F2'}), ('493-27-00631', {})],
    [('계약 종료일', {'bold':True,'center':True,'bg':'F2F2F2'}), ('2025년 9월 30일', {})],
    [('대조 기간', {'bold':True,'center':True,'bg':'F2F2F2'}), ('2021년 4월 ~ 2025년 9월 (계약 전 기간)', {})],
]
make_table(doc, ['항목', '내용'], info_rows, [4.0, 11.5])

# ===== 3. ERP ACC 외상매출금 수주 대조 =====
add_heading(doc, '3. ERP ACC 외상매출금 수주 대조')
add_body(doc, '대조 기간: 2021년 4월 ~ 2025년 10월  (단위: 원, VAT 포함)')

erp_rows = [
    [('외상매출금 (a)', {'bold':True,'center':True,'bg':'F2F2F2'}),
     ('7,186,740,259', {'center':True})],
    [('외상매출금 입금 (b)', {'bold':True,'center':True,'bg':'F2F2F2'}),
     ('7,037,963,889', {'center':True})],
    [('SPC 결제수수료 (c)', {'bold':True,'center':True,'bg':'F2F2F2'}),
     ('149,047,373', {'center':True})],
    [('반품수수료 잡이익 (d)', {'bold':True,'center':True,'bg':'F2F2F2'}),
     ('251,000', {'center':True})],
    [('진여채권 (a-b-c+d)', {'bold':True,'center':True,'bg':'FFF0F0','fc':(180,0,0)}),
     ('△20,003', {'bold':True,'center':True,'bg':'FFF0F0','fc':(180,0,0)})],
]
make_table(doc, ['구분', '금액'], erp_rows, [7.5, 8.0])

add_note(doc, '※ 진여채권 산정: 7,186,740,259 - 7,037,963,889 - 149,047,373 + 251,000 = △20,003원')
add_note(doc, '※ 외상매출금 중 2026-06 24,646,000원은 롯데구리점서브 오결제 이관 건 (거래처 이관 별도 처리, 하단 특이사항 참조)')

# ===== 4. PG사 청구내역 vs ERP 대조 =====
add_heading(doc, '4. PG사 청구내역 vs ERP 대조 (연도별)')

pg_rows = [
    [('2021', {'center':True}), ('1,202,777,551', {'center':True}), ('1,178,253,148', {'center':True}), ('24,524,403', {'center':True})],
    [('2022', {'center':True}), ('1,546,909,261', {'center':True}), ('1,515,795,249', {'center':True}), ('31,114,012', {'center':True})],
    [('2023', {'center':True}), ('1,599,294,047', {'center':True}), ('1,565,444,315', {'center':True}), ('33,849,732', {'center':True})],
    [('2024', {'center':True}), ('1,594,859,930', {'center':True}), ('1,562,509,549', {'center':True}), ('32,350,381', {'center':True})],
    [('2025', {'center':True}), ('1,243,170,473', {'center':True}), ('1,215,961,628', {'center':True}), ('27,208,845', {'center':True})],
    [('합계', {'bold':True,'center':True,'bg':'F2F2F2'}), ('7,187,011,262', {'bold':True,'center':True,'bg':'F2F2F2'}),
     ('7,037,963,889', {'bold':True,'center':True,'bg':'F2F2F2'}), ('149,047,373', {'bold':True,'center':True,'bg':'F2F2F2'})],
    [('ERP 외상매출금', {'bold':True,'center':True,'bg':'F2F2F2'}), ('7,186,740,259', {'center':True,'bg':'F2F2F2'}), ('—', {'center':True,'bg':'F2F2F2'}), ('—', {'center':True,'bg':'F2F2F2'})],
    [('차이 (ERP-PG)', {'bold':True,'center':True,'bg':'FFF0F0','fc':(180,0,0)}),
     ('△271,003', {'bold':True,'center':True,'bg':'FFF0F0','fc':(180,0,0)}),
     ('—', {'center':True,'bg':'FFF0F0'}), ('→ 잡이익 처리', {'center':True,'bg':'FFF0F0'})],
]
make_table(doc, ['연도', '정산대상액 (PG 청구)', '실지급액', '결제수수료'], pg_rows, [2.5, 4.5, 4.5, 4.0])

# ===== 5. 잔여채권 정산 처리 내역 =====
add_heading(doc, '5. 잔여채권 / 미수금 정산 처리')

settle_rows = [
    [('PG사 환급금\n(역환)', {'center':True}),
     ('외상매출금1', {'center':True}), ('0', {'center':True}), ('0', {'center':True}),
     ('LXH목동', {'center':True}), ('역환금액 없음', {'center':True})],
    [('', {'center':True}),
     ('보통예금', {'center':True}), ('0', {'center':True}), ('0', {'center':True}),
     ('섹타나인', {'center':True}), ('', {'center':True})],
    [('잔여채권 정리', {'bold':True,'center':True,'bg':'FFF0F0'}),
     ('외상매출금1', {'center':True,'bg':'FFF0F0'}), ('20,003', {'bold':True,'center':True,'bg':'FFF0F0','fc':(180,0,0)}), ('0', {'center':True,'bg':'FFF0F0'}),
     ('LXH목동', {'center':True,'bg':'FFF0F0'}), ('잡이익 처리', {'bold':True,'center':True,'bg':'FFF0F0'})],
    [('', {'center':True}),
     ('잡이익', {'center':True}), ('0', {'center':True}), ('20,003', {'bold':True,'center':True,'fc':(180,0,0)}),
     ('', {'center':True}), ('', {'center':True})],
]
make_table(doc, ['구분', '계정과목', '차변 (원)', '대변 (원)', '거래처', '비고'], settle_rows,
           [2.8, 2.5, 2.0, 2.0, 2.5, 3.7])

add_note(doc, '※ 해당 보고서 및 전표를 통하여 신도림3(구 LXH목동 계정)에 대한 최종 채권 정리 및 코드 폐쇄 예정')

# ===== 6. 특이사항 — 롯데구리점서브 오결제 이관 =====
add_heading(doc, '6. 특이사항 — 롯데구리점서브 오결제 이관 건')
add_body(doc, '롯데구리점서브 전시품 수주 결제건(2025-05)이 LXH목동 단말기로 처리되어,')
add_body(doc, '해당 결제금액이 LXH목동 계정에 귀속된 건으로, 별도 거래처 이관 정정 처리함.')

ogul_rows = [
    [('이관 금액', {'bold':True,'center':True,'bg':'F2F2F2'}), ('24,646,000원 (5건)', {'center':True})],
    [('원인', {'bold':True,'center':True,'bg':'F2F2F2'}), ('동일 점주 복수 매장 운영 — LXH목동 단말기로 롯데구리 수주 결제', {})],
    [('처리 방법', {'bold':True,'center':True,'bg':'F2F2F2'}), ('거래처 이관 정정 (LXH목동 → 롯데구리점서브)', {})],
    [('처리 시기', {'bold':True,'center':True,'bg':'F2F2F2'}), ('2026년 6월 (Connect+ 시스템 이관 의뢰)', {})],
    [('비고', {'bold':True,'center':True,'bg':'F2F2F2'}), ('채권 정산보고서 대조 범위 외 항목으로 별도 관리', {})],
]
make_table(doc, ['항목', '내용'], ogul_rows, [4.0, 11.5])

# ===== 7. 요청사항 =====
add_heading(doc, '7. 요청사항')
add_body(doc, '1) 잔여채권(△20,003원) 잡이익 처리 전표 승인 → 일룸 재경팀')
add_body(doc, '2) 신도림3 (구 LXH목동 계정, 61LM06) 코드 폐쇄 처리 → 일룸 재경팀')
add_note(doc, '   • 계약 종료 매장으로 추가 수주 없음, 잔여 채권 정리 완료 후 계정 폐쇄 요청')

# ===== 8. 첨부자료 =====
add_heading(doc, '8. 첨부자료')
add_body(doc, '1) (채권대사) LXH목동점 V.2 — ERP ACC / 청구내역 / 지급내역 / 미결초과 상세')
add_body(doc, '2) 거래처별보조원장_LXH목동(61LM96)_260615')
add_body(doc, '3) 거래처별원장_LXH목동(61LM96)_260615')
add_body(doc, '4) (채권대사) 롯데구리점서브 오결제 이관 정산 관련 자료')

# 저장
out = r'c:\Users\FURSYS\Downloads\iloom-workspace-claude\20-operations\23-채산관리\2025_LXH목동점_채권정산보고서.docx'
doc.save(out)
print(f'saved: {out}')
