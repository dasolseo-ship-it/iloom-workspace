from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# 페이지 여백
section = doc.sections[0]
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(3.0)
section.right_margin = Cm(2.5)

# 기본 폰트 설정
style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style.font.size = Pt(10)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

def set_font(run, size=10, bold=False):
    run.font.name = '맑은 고딕'
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

def add_para(text, align=WD_ALIGN_PARAGRAPH.LEFT, size=10, bold=False, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_font(run, size, bold)
    return p

def set_cell_border(cell, top=True, bottom=True, left=True, right=True):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, flag in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if flag:
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tcBorders.append(border)
    tcPr.append(tcBorders)

def cell_text(cell, text, bold=False, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    set_font(run, size, bold)

def shade_cell(cell, hex_color='D9D9D9'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

# 제목
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(12)
run = p.add_run('2026년 5월 일룸 송도점 전대료 정산')
set_font(run, 14, True)

# 도입문
add_para('아래와 같이 송도점 전대료를 정산하고자 하오니 재가 바랍니다.',
         align=WD_ALIGN_PARAGRAPH.CENTER, size=10, space_after=4)
add_para('----- 아   래 -----',
         align=WD_ALIGN_PARAGRAPH.CENTER, size=10, space_after=12)

# 1. 개요
add_para('1. 개요', bold=True, size=10, space_before=4, space_after=2)
items = [
    '1) 일룸 송도점은 2025년 9월 1일부터 2026년 8월 31일까지 재계약(1년 연장)을 진행하였으며, 2025년 9월 1일부터 2026년 8월 31일까지 홈플러스 기업회생절차 리스크를 최소화하기 위해 기존 홈플러스 단말기 결제방식에서 일룸 단말기 결제방식으로 변경함',
    '2) 이에 따라 2025년 9월부터 후정산 형태로 송도점의 월별 전대료(수수료율)를 정산하고자 함',
    '3) 월별 전대료는 영업 시작 해당 월 1일부터 해당 월 말일을 기준으로 영업기간 중 발생된 매출 분에 대해 익월 15일 전까지 수수료를 전대인에게 입금하기로 함 (계약서 제 8조의 2에 의거)',
    '4) 홈플러스 측 요청에 의해 2026년 5월 1일부터 2026년 6월 30일까지 계약조건(전대료율)이 한시적으로 변경됨\n     - 변경 사유 : 37개점 하이퍼 휴점으로 인한 한시적 임대 조건 완화 (상호 합의)\n     - 협의 일시 : 2026-05-29  /  참석자 : 홈플러스 최윤도 과장, 전차인 정보은 대표\n     - 계약번호 : 2509920',
]
for item in items:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(item)
    set_font(run, 9.5)

# 2. 계약조건 변경
add_para('2. 계약조건 변경 사항', bold=True, size=10, space_before=8, space_after=2)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(2)
run = p.add_run('(단위 : 원 / VAT 별도)')
set_font(run, 9)

tbl = doc.add_table(rows=2, cols=3)
tbl.style = 'Table Grid'
headers = ['변경조건 적용기간', '기존 계약조건', '변경 계약조건']
for i, h in enumerate(headers):
    cell_text(tbl.rows[0].cells[i], h, bold=True, size=9.5)
    shade_cell(tbl.rows[0].cells[i])

data_row = tbl.rows[1]
cell_text(data_row.cells[0], '2026-05-01 ~ 2026-06-30', size=9.5)
cell_text(data_row.cells[1],
    '전대료(수수료율) :\n· 월 순매출의 8.2%\n· 월 순매출 90,000,000원 미만 시\n  정액 7,380,000원 (부가세 별도)\n· 월 순매출 150,000,000원 초과분\n  수수료율 6.2% 적용',
    size=9.5, align=WD_ALIGN_PARAGRAPH.LEFT)
cell_text(data_row.cells[2],
    '전대료(수수료율) :\n· 월 순매출의 7.4%\n· 월 순매출 90,000,000원 미만 시\n  정액 7,380,000원 (부가세 별도)\n· 월 순매출 150,000,000원 초과분\n  수수료율 5.5% 적용',
    size=9.5, align=WD_ALIGN_PARAGRAPH.LEFT)

for row in tbl.rows:
    for cell in row.cells:
        set_cell_border(cell)

# 3. 매장현황
add_para('3. 매장현황', bold=True, size=10, space_before=10, space_after=2)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(2)
run = p.add_run('(단위 : 원 / VAT 별도)')
set_font(run, 9)

tbl2 = doc.add_table(rows=5, cols=3)
tbl2.style = 'Table Grid'
col_w = [Cm(2.0), Cm(9.0), Cm(4.5)]
for row in tbl2.rows:
    for i, cell in enumerate(row.cells):
        cell.width = col_w[i]

headers2 = ['구  분', '내  용', '비  고']
for i, h in enumerate(headers2):
    cell_text(tbl2.rows[0].cells[i], h, bold=True, size=9.5)
    shade_cell(tbl2.rows[0].cells[i])

rows_data = [
    ('매 장 명', '송도점', ''),
    ('주  소', '홈플러스 송도점 (인천광역시 연수구 송도국제대로 165), 1층 X7,YL 코너', ''),
    ('면  적', '444.2㎡ (134.4평)', ''),
    ('전 대 료', '월 순매출액의 7.4%\n(월 순매출 90,000,000원 미만 시 정액 7,380,000원 적용)', '2026.05.01~06.30 적용\n익월 15일 지급'),
]
for i, (a, b, c) in enumerate(rows_data):
    row = tbl2.rows[i+1]
    cell_text(row.cells[0], a, size=9.5)
    cell_text(row.cells[1], b, size=9.5, align=WD_ALIGN_PARAGRAPH.LEFT)
    cell_text(row.cells[2], c, size=9.5)

for row in tbl2.rows:
    for cell in row.cells:
        set_cell_border(cell)

# 4. 정산현황
add_para('4. 정산현황 (2026년 5월 1일 ~ 5월 31일)', bold=True, size=10, space_before=10, space_after=2)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(2)
run = p.add_run('(단위 : 원 / VAT 포함)')
set_font(run, 9)

add_para('① 섹터나인 결제 내역', size=9.5, space_before=2, space_after=2)
tbl3 = doc.add_table(rows=4, cols=3)
tbl3.style = 'Table Grid'
for i, h in enumerate(['결제수단', '건수', '정산대상금액']):
    cell_text(tbl3.rows[0].cells[i], h, bold=True, size=9.5)
    shade_cell(tbl3.rows[0].cells[i])
s_data = [('가상계좌', '8', '7,531,000'), ('신용카드', '44', '29,187,500'), ('합  계', '52', '36,718,500')]
for i, (a, b, c) in enumerate(s_data):
    bold = (i == 2)
    cell_text(tbl3.rows[i+1].cells[0], a, bold=bold, size=9.5)
    cell_text(tbl3.rows[i+1].cells[1], b, bold=bold, size=9.5)
    cell_text(tbl3.rows[i+1].cells[2], c, bold=bold, size=9.5)
for row in tbl3.rows:
    for cell in row.cells:
        set_cell_border(cell)

add_para('② 홈플러스 결제 내역', size=9.5, space_before=6, space_after=2)
tbl4 = doc.add_table(rows=2, cols=4)
tbl4.style = 'Table Grid'
for i, h in enumerate(['단말기', '매출', '수수료', '비고']):
    cell_text(tbl4.rows[0].cells[i], h, bold=True, size=9.5)
    shade_cell(tbl4.rows[0].cells[i])
cell_text(tbl4.rows[1].cells[0], '일룸 POS', size=9.5)
cell_text(tbl4.rows[1].cells[1], '37,567,500', size=9.5)
cell_text(tbl4.rows[1].cells[2], '3,080,535', size=9.5)
cell_text(tbl4.rows[1].cells[3], '일룸이 지급해야 할 후정산 수수료', size=9.5, align=WD_ALIGN_PARAGRAPH.LEFT)
for row in tbl4.rows:
    for cell in row.cells:
        set_cell_border(cell)

add_para('③ 5월 확정 매출', size=9.5, space_before=6, space_after=2)
tbl5 = doc.add_table(rows=4, cols=2)
tbl5.style = 'Table Grid'
for i, h in enumerate(['항목', '금액']):
    cell_text(tbl5.rows[0].cells[i], h, bold=True, size=9.5)
    shade_cell(tbl5.rows[0].cells[i])
m_data = [
    ('카드사 결제 합계', '81,594,300'),
    ('카드사 취소 (2026.06.01 처리, 1,576,000원 × 2건)', '△ 3,152,000'),
    ('5월 확정 매출', '78,442,300'),
]
for i, (a, b) in enumerate(m_data):
    bold = (i == 2)
    cell_text(tbl5.rows[i+1].cells[0], a, bold=bold, size=9.5, align=WD_ALIGN_PARAGRAPH.LEFT)
    cell_text(tbl5.rows[i+1].cells[1], b, bold=bold, size=9.5)
    if bold:
        shade_cell(tbl5.rows[i+1].cells[0], 'F2F2F2')
        shade_cell(tbl5.rows[i+1].cells[1], 'F2F2F2')
for row in tbl5.rows:
    for cell in row.cells:
        set_cell_border(cell)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(3)
p.paragraph_format.space_after = Pt(2)
run = p.add_run('※ 취소 2건(3,152,000원)은 카드사에서 2026.06.01 취소 처리됨. 해당 금액은 6월 매출에서 차감 적용 예정.')
set_font(run, 9)

# 5. 전대료 산출
add_para('5. 전대료 산출', bold=True, size=10, space_before=10, space_after=2)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(2)
run = p.add_run('(단위 : 원 / VAT 별도)')
set_font(run, 9)

notes = [
    '· 5월 확정 매출 (VAT 포함) : 78,442,300원',
    '· 5월 확정 매출 (VAT 별도) : 약 71,311,000원  →  90,000,000원 미만  →  정액 적용',
]
for n in notes:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(n)
    set_font(run, 9.5)

tbl6 = doc.add_table(rows=4, cols=2)
tbl6.style = 'Table Grid'
for i, h in enumerate(['구  분', '금  액']):
    cell_text(tbl6.rows[0].cells[i], h, bold=True, size=9.5)
    shade_cell(tbl6.rows[0].cells[i])
fee_data = [
    ('전대료 (부가세 별도)', '7,380,000'),
    ('부가세 (10%)', '738,000'),
    ('전대료 합계 (부가세 포함)', '8,118,000'),
]
for i, (a, b) in enumerate(fee_data):
    bold = (i == 2)
    cell_text(tbl6.rows[i+1].cells[0], a, bold=bold, size=9.5)
    cell_text(tbl6.rows[i+1].cells[1], b, bold=bold, size=9.5)
    if bold:
        shade_cell(tbl6.rows[i+1].cells[0], 'F2F2F2')
        shade_cell(tbl6.rows[i+1].cells[1], 'F2F2F2')
for row in tbl6.rows:
    for cell in row.cells:
        set_cell_border(cell)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(3)
p.paragraph_format.space_after = Pt(2)
run = p.add_run('※ 월 순매출 90,000,000원 미만으로 정액 7,380,000원 적용 (기존·변경 조건 동일)')
set_font(run, 9)
p2 = doc.add_paragraph()
p2.paragraph_format.space_before = Pt(0)
p2.paragraph_format.space_after = Pt(2)
run2 = p2.add_run('※ 납부 기한 : 세금계산서 수령 후 익월 15일 18:00까지')
set_font(run2, 9)

# 6. 첨부
add_para('6. 첨부', bold=True, size=10, space_before=10, space_after=2)
for att in ['1) 계약조건 변경 합의서 (계약번호 2509920, 협의일 2026-05-29)',
            '2) 홈플러스 ECS 정산서 (2026년 5월)']:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(att)
    set_font(run, 9.5)

add_para('- 이  상 -', align=WD_ALIGN_PARAGRAPH.CENTER, size=10, space_before=16)

out = r'c:\Users\FURSYS\Downloads\iloom-workspace-claude\20-operations\23-채산관리\2026-05_송도점_전대료_정산보고서.docx'
doc.save(out)
print(f'저장완료: {out}')
