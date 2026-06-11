# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt

REF = r'C:\Users\FURSYS\Downloads\일룸_하절기 연장영업 미참여 요청 공문_신세계시흥.docx'
OUT = r'C:\Users\FURSYS\Downloads\iloom-workspace-claude\00-inbox\2026-06-11_일룸송도점_운영종료_퇴점통보_홈플러스_v2.docx'

doc = Document(REF)
ref_table = doc.tables[0]

def clear_cell(cell):
    tc = cell._tc
    for p in tc.findall(qn('w:p')):
        tc.remove(p)
    new_p = OxmlElement('w:p')
    tc.append(new_p)

def set_cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    clear_cell(cell)
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold

# Row 0: 발신 주소
set_cell_text(
    ref_table.rows[0].cells[2],
    '발신     (05661) 서울시 송파구 오금로 311 퍼시스빌딩 8층  ㈜일룸'
)

# Row 1: 문서번호
set_cell_text(ref_table.rows[1].cells[2], '리테일사업팀-제260611-1호')

# Row 2: 수신
set_cell_text(ref_table.rows[2].cells[2], '홈플러스 주식회사  Mall사업부문  계약 담당자')

# Row 3: 작성일
set_cell_text(ref_table.rows[3].cells[2], '2026년 06월 11일')

# Row 4: 담당자
set_cell_text(ref_table.rows[4].cells[2], '일룸 리테일사업팀')

# Row 5: 제목
set_cell_text(
    ref_table.rows[5].cells[2],
    '일룸 C)인천송도점 운영 종료(퇴점) 통보의 건',
    bold=True
)

# ─── 본문 셀 (Row 6) ────────────────────────────────────────────────────────────
body_cell = ref_table.rows[6].cells[0]
tc = body_cell._tc
for p_el in list(tc.findall(qn('w:p'))):
    tc.remove(p_el)

def new_para(text='', align=WD_ALIGN_PARAGRAPH.LEFT, bold=False, before=0, after=60):
    p_el = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    jc = OxmlElement('w:jc')
    jc_map = {
        WD_ALIGN_PARAGRAPH.LEFT: 'left',
        WD_ALIGN_PARAGRAPH.CENTER: 'center',
        WD_ALIGN_PARAGRAPH.RIGHT: 'right',
    }
    jc.set(qn('w:val'), jc_map.get(align, 'left'))
    pPr.append(jc)
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(int(before)))
    spacing.set(qn('w:after'), str(int(after)))
    pPr.append(spacing)
    p_el.append(pPr)
    if text:
        r_el = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        if bold:
            b_el = OxmlElement('w:b')
            rPr.append(b_el)
        r_el.append(rPr)
        t = OxmlElement('w:t')
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t.text = text
        r_el.append(t)
        p_el.append(r_el)
    return p_el

def add_p(text='', align=WD_ALIGN_PARAGRAPH.LEFT, bold=False, before=0, after=60):
    tc.append(new_para(text, align, bold, before, after))

add_p('1.  귀사의 무궁한 발전을 기원합니다.', after=80)
add_p()
add_p('2.  당사는 귀사와의 Mall 전대차 계약(계약번호: 2509920)에 의거하여 홈플러스 C)인천송도점 1층 X7·YL 코너에서 \'일룸 C)인천송도점\'을 운영하고 있으며, 그간 보내주신 협조와 지원에 깊이 감사드립니다.', after=80)
add_p()
add_p('3.  당사는 현행 전대차 계약의 만료일(2026년 08월 31일)에 따라 운영을 종료하고 퇴점할 예정임을 아래와 같이 공식 통보합니다.', after=80)
add_p()
add_p('- 다  음 -', align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, before=40, after=40)
add_p()
add_p('가.  매 장 명        :  일룸 C)인천송도점', after=60)
add_p('나.  계 약 번 호  :  2509920', after=60)
add_p('다.  소 재 지        :  인천광역시 연수구 송도국제대로 165  홈플러스 C)인천송도점 1층 X7·YL 코너', after=60)
add_p('라.  계 약 기 간  :  2025년 09월 01일 ~ 2026년 08월 31일', after=60)
add_p('마.  운영 종료(퇴점) 예정일  :  2026년 08월 31일', after=60)
add_p()
add_p('퇴점 관련 아래 사항에 대해 협조를 부탁드리며, 추가 협의가 필요한 경우 담당자(일룸 리테일 사업팀)에게 연락 주시기 바랍니다. 오랜 기간 협력해 주신 점 거듭 감사드립니다.', after=80)
add_p()
add_p('바.  전대보증금(금 33,325,000원) 반환  :  계약 종료 및 매장 인도에 따른 전대보증금 반환 절차 안내 요청', after=60)
add_p('사.  잔여 행정 절차  :  전산 처리, POS 기기 반환 등 퇴점 관련 행정 절차 협조 요청', after=60)
add_p()
add_p()
add_p()
add_p('㈜일룸      대표이사  정 보 은',
      align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, before=80, after=40)
add_p()

# ─── 수신처 행 (Row 7) ──────────────────────────────────────────────────────────
set_cell_text(ref_table.rows[7].cells[1],
              '홈플러스 주식회사  Mall사업부문  계약 담당자\n(서울특별시 강서구 화곡로 398  홈플러스 본사)')

doc.save(OUT)
print('saved:', OUT)
